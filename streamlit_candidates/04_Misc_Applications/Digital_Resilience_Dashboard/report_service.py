import streamlit as st
st.title('report_service.py - 自動化展示')
st.info('這是從原專案腳本自動包裝產生的互動式介面。')

# -*- coding: utf-8 -*-
"""
report_service.py — 自動化稽核報告生成器 (一鍵匯出 Word)
功能：
  - 實作 generate_audit_word_report 函數，渲染符合 ISO 規範的內部稽核與 AI 系統運行週報。
  - 採用 python-docx 進行流式文件處理，支援邊距設定、中文字型定義、自訂表格背景色與內距設定。
"""
import io
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex: str):
    """設定儲存格背景顏色 (Shading)"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top: int = 120, bottom: int = 120, left: int = 180, right: int = 180):
    """設定儲存格內邊距 (Padding)，數值單位為 dxa (1 pt = 20 dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin_name)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_paragraph(doc, text: str, style_name: str = 'Normal', space_after: int = 6, 
                         bold: bool = False, size_pt: float = 10.5, color_rgb = None) -> OxmlElement:
    """新增具備指定中文字型與樣式的段落"""
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = 'Microsoft JhengHei'
    if color_rgb:
        run.font.color.rgb = color_rgb
        
    # 強制支援 Word 內之中文字型渲染 (Microsoft JhengHei / 微軟正黑體)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    rPr.append(rFonts)
    return p

def add_cell_text(cell, text: str, bold: bool = False, size_pt: float = 9.5, 
                  color_rgb = None, alignment = WD_ALIGN_PARAGRAPH.LEFT):
    """向儲存格寫入格式化文字，支援微軟正黑體與換行符號"""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    # 支援換行寫入
    lines = text.split('\n')
    for idx, line in enumerate(lines):
        if idx > 0:
            p.add_run('\n')
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.name = 'Microsoft JhengHei'
        if color_rgb:
            run.font.color.rgb = color_rgb
            
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
        rPr.append(rFonts)

def generate_audit_word_report(logs_data: list, ai_summary: str) -> bytes:
    """
    動態生成 Word 查核報表，並將結果輸出為 bytes 流。
    logs_data: 過去 7 日的審計日誌列表
    ai_summary: LLM 主導稽核員生成的本週合規摘要
    """
    doc = Document()
    
    # 1. 頁面邊距設定 (標準 A4 邊距，上下左右 1 英吋)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # 色彩定義 (Corporate Theme)
    COLOR_PRIMARY = RGBColor(27, 54, 93)      # 科技深藍 #1B365D
    COLOR_SECONDARY = RGBColor(111, 126, 150) # 灰藍色
    COLOR_WHITE = RGBColor(255, 255, 255)
    
    # 2. 報告標題 (Title)
    add_styled_paragraph(
        doc, 
        "企業內部稽核與 AI 系統運行綜合報告", 
        space_after=4, 
        bold=True, 
        size_pt=18.0, 
        color_rgb=COLOR_PRIMARY
    )
    add_styled_paragraph(
        doc, 
        "ISO 27001 資訊安全與 ISO 14064-1 溫室氣體盤查合規運行軌跡", 
        space_after=18, 
        bold=False, 
        size_pt=12.0, 
        color_rgb=COLOR_SECONDARY
    )
    
    # 3. 報告屬性 (Metadata Table)
    today_str = datetime.date.today().strftime("%Y-%m-%d")
    date_7_days_ago = (datetime.date.today() - datetime.timedelta(days=7)).strftime("%Y-%m-%d")
    
    # 建立 2 行 4 列的輕量屬性表格
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.autofit = False
    
    # 設定寬度
    widths = [Inches(1.2), Inches(2.0), Inches(1.2), Inches(2.1)]
    for row in meta_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            
    # 寫入屬性欄位
    add_cell_text(meta_table.cell(0, 0), "報告日期", bold=True, size_pt=9.5, color_rgb=COLOR_PRIMARY)
    add_cell_text(meta_table.cell(0, 1), today_str, size_pt=9.5)
    add_cell_text(meta_table.cell(0, 2), "報告區間", bold=True, size_pt=9.5, color_rgb=COLOR_PRIMARY)
    add_cell_text(meta_table.cell(0, 3), f"{date_7_days_ago} 至 {today_str}", size_pt=9.5)
    
    add_cell_text(meta_table.cell(1, 0), "製表系統", bold=True, size_pt=9.5, color_rgb=COLOR_PRIMARY)
    add_cell_text(meta_table.cell(1, 1), "企業數位韌性 AI 導航系統", size_pt=9.5)
    add_cell_text(meta_table.cell(1, 2), "製表單位", bold=True, size_pt=9.5, color_rgb=COLOR_PRIMARY)
    add_cell_text(meta_table.cell(1, 3), "企業數位安全與 ESG 稽核中樞", size_pt=9.5)
    
    # 為屬性表美化背景與邊距
    for row in meta_table.rows:
        for cell in row.cells:
            set_cell_background(cell, "F4F6F9")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12) # 間隔行
    
    # 4. 高階主管摘要 (Executive Summary)
    add_styled_paragraph(doc, "一、 高階主管摘要 (Executive Summary)", bold=True, size_pt=14.0, color_rgb=COLOR_PRIMARY)
    
    # 建立一個有左側藍色裝飾框的效果段落
    summary_box = doc.add_table(rows=1, cols=1)
    summary_box.autofit = False
    summary_box.columns[0].width = Inches(6.5)
    cell = summary_box.cell(0, 0)
    set_cell_background(cell, "F0F4F8") # 淺藍背景
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # 寫入 AI 摘要內容
    summary_text = ai_summary.strip() if ai_summary else "本週系統運作良好，無任何安全或環境合規異常事件記錄。"
    add_cell_text(cell, summary_text, size_pt=10.0)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12) # 間隔行
    
    # 5. 稽核軌跡詳情 (Audit Trail Details)
    add_styled_paragraph(doc, "二、 稽核軌跡詳情 (Audit Trail Details)", bold=True, size_pt=14.0, color_rgb=COLOR_PRIMARY)
    add_styled_paragraph(
        doc, 
        "以下為近 7 天內系統記錄的使用者操作事件、安全防禦動作及 AI 對策評估：", 
        space_after=8, 
        size_pt=10.0
    )
    
    # 建立稽核事件資料表 (Table Grid)
    # 欄位：時間、事件/動作、AI 建議對策、狀態
    cols_num = 4
    if not logs_data:
        # 空報表防呆處理
        table = doc.add_table(rows=2, cols=cols_num)
        table.autofit = False
        table.style = 'Table Grid'
        
        # 寬度分配
        col_widths = [Inches(1.3), Inches(2.2), Inches(2.2), Inches(0.8)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
                
        # 寫入標頭
        headers = ["時間", "事件/動作", "AI 建議對策", "狀態"]
        for idx, text in enumerate(headers):
            cell = table.cell(0, idx)
            add_cell_text(cell, text, bold=True, size_pt=9.5, color_rgb=COLOR_WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_background(cell, "1B365D")  # 深藍標頭背景
            set_cell_margins(cell, top=140, bottom=140, left=100, right=100)
            
        # 寫入空提示行 (合併中間)
        cell_span = table.cell(1, 0)
        for col_idx in range(1, cols_num):
            cell_span.merge(table.cell(1, col_idx))
        add_cell_text(cell_span, "【防呆提示】本週無異常事件與稽核軌跡記錄。", size_pt=9.5, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_margins(cell_span, top=180, bottom=180, left=100, right=100)
        
    else:
        # 有稽核日誌，渲染資料表格
        table = doc.add_table(rows=len(logs_data) + 1, cols=cols_num)
        table.autofit = False
        table.style = 'Table Grid'
        
        col_widths = [Inches(1.3), Inches(2.2), Inches(2.2), Inches(0.8)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
                
        # 1. 寫入標頭
        headers = ["時間", "事件/動作", "AI 建議對策", "狀態"]
        for idx, text in enumerate(headers):
            cell = table.cell(0, idx)
            add_cell_text(cell, text, bold=True, size_pt=9.5, color_rgb=COLOR_WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_background(cell, "1B365D")  # 深藍背景
            set_cell_margins(cell, top=140, bottom=140, left=100, right=100)
            
        # 2. 寫入日誌內容
        for r_idx, r in enumerate(logs_data, start=1):
            row = table.rows[r_idx]
            
            # 時間格式整理 (取至秒)
            raw_ts = r.get("timestamp", "")
            ts_str = raw_ts[:19].replace("T", " ") if raw_ts else "未知時間"
            
            # 事件動作拼裝
            user = r.get("username", "System")
            role = r.get("role", "SYSTEM")
            details = r.get("action_details", "")
            prompt = r.get("prompt", "")
            
            event_text = f"【{user} ({role})】\n{details}"
            if prompt:
                event_text += f"\n[輸入: {prompt}]"
                
            # AI 建議對策
            ai_res = r.get("ai_response", "")
            if not ai_res:
                if r.get("action_type") == "Approval":
                    ai_res = "已由管理員人工審核通過，執行對應通報。"
                elif r.get("action_type") == "Security_Block":
                    ai_res = "安全引擎已攔截此操作並阻斷發布。"
                else:
                    ai_res = "無/不適用 (非 AI 推理事件)"
            
            # 狀態顯示與標識
            act_type = r.get("action_type", "")
            if act_type == "Security_Block":
                status_text = "已攔截"
                status_color = RGBColor(220, 38, 38) # 紅色
            elif act_type == "Approval":
                status_text = "已核准"
                status_color = RGBColor(22, 163, 74) # 綠色
            else:
                status_text = "已記錄"
                status_color = RGBColor(30, 41, 59)  # 深灰黑色
                
            # 寫入各單元格
            add_cell_text(row.cells[0], ts_str, size_pt=9.0)
            add_cell_text(row.cells[1], event_text, size_pt=9.0)
            add_cell_text(row.cells[2], ai_res, size_pt=9.0)
            add_cell_text(row.cells[3], status_text, bold=True, size_pt=9.0, color_rgb=status_color, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # 交替斑馬紋背景色與 Padding
            bg_color = "FFFFFF" if r_idx % 2 != 0 else "F9FAFB"
            for cell in row.cells:
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
                
    # 6. 文件結尾簽章
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_footer = p_footer.add_run("報告審核人：____________________ (簽章)\n報告核准人：____________________ (簽章)")
    run_footer.font.name = 'Microsoft JhengHei'
    run_footer.font.size = Pt(10.0)
    
    rPr = run_footer._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    rPr.append(rFonts)
    
    # 7. 寫出至記憶體 BytesIO 流
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue()


# ====== 自動生成的測試執行區塊 ======
if __name__ == "__main__":
    st.write("---")
    st.subheader("函數測試區塊")
    if st.button("執行 set_cell_background"):
        try:
            res = set_cell_background() # 請視需要自行補上參數
            st.write("執行結果:", res)
        except Exception as e:
            st.error(f"執行出錯，可能需要提供參數：{e}")
    if st.button("執行 set_cell_margins"):
        try:
            res = set_cell_margins() # 請視需要自行補上參數
            st.write("執行結果:", res)
        except Exception as e:
            st.error(f"執行出錯，可能需要提供參數：{e}")
    if st.button("執行 add_styled_paragraph"):
        try:
            res = add_styled_paragraph() # 請視需要自行補上參數
            st.write("執行結果:", res)
        except Exception as e:
            st.error(f"執行出錯，可能需要提供參數：{e}")
    if st.button("執行 add_cell_text"):
        try:
            res = add_cell_text() # 請視需要自行補上參數
            st.write("執行結果:", res)
        except Exception as e:
            st.error(f"執行出錯，可能需要提供參數：{e}")
    if st.button("執行 generate_audit_word_report"):
        try:
            res = generate_audit_word_report() # 請視需要自行補上參數
            st.write("執行結果:", res)
        except Exception as e:
            st.error(f"執行出錯，可能需要提供參數：{e}")
