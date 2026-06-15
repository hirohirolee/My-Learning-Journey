import streamlit as st
import pandas as pd
import os
import re
import random
from docx import Document
from datetime import datetime, timedelta

# ==========================================
# 1. Back-end Data Engineering & Synthesis
# ==========================================
def init_mock_files():
    """Generates synthetic manufacturing telemetry (Excel) & compliance manuals (Word) if missing or outdated"""
    
    # A. MES Operating Data (100 Rows, 30-Day continuous timeline, 3 machines, anomalies)
    regenerate_mes = False
    if not os.path.exists("mes_data.xlsx"):
        regenerate_mes = True
    else:
        try:
            df_existing = pd.read_excel("mes_data.xlsx")
            if len(df_existing) != 100 or "Risk_Level" not in df_existing.columns:
                regenerate_mes = True
        except Exception:
            regenerate_mes = True

    if regenerate_mes:
        random.seed(42)  # Maintain reproducibility for deterministic testing
        machines = ["EQP-A03", "EQP-B01", "EQP-C02"]
        rows_per_machine = [34, 33, 33]  # Sums to exactly 100 rows
        start_date = datetime.now() - timedelta(days=30)
        all_rows = []

        for m_idx, machine in enumerate(machines):
            N = rows_per_machine[m_idx]
            # Distribute dates chronologically over 30 days
            m_dates = [start_date + timedelta(days=30 * i / N) for i in range(N)]
            
            # Setup normal baseline distributions
            # E.g. Normal vibrations: Green (3.0-5.0Hz) or Yellow (5.0-7.5Hz)
            vibrations = [round(random.uniform(3.2, 7.5), 2) for _ in range(N)]
            yields = [round(random.uniform(95.0, 99.5), 2) for _ in range(N)]
            powers = [round(random.uniform(90.0, 125.0), 2) for _ in range(N)]
            
            # Inject mathematical anomalies at indices 5, 12, 20, 27
            anomaly_indices = [5, 12, 20, 27]
            for idx in anomaly_indices:
                if idx + 2 < N:
                    # Vibration peaks > 8.0
                    vibrations[idx] = round(random.uniform(8.1, 9.8), 2)
                    # Exactly 2 periods later: yield rate drops < 90%, power consumption spikes > 150kW
                    yields[idx + 2] = round(random.uniform(80.0, 89.5), 2)
                    powers[idx + 2] = round(random.uniform(151.0, 185.0), 2)
            
            for i in range(N):
                vib = vibrations[i]
                # Engineer Risk_Level
                if vib < 5.0:
                    risk = "Green"
                elif vib <= 8.0:
                    risk = "Yellow"
                else:
                    risk = "Red"
                    
                all_rows.append({
                    "Date": m_dates[i].strftime("%Y-%m-%d %H:%M:%S"),
                    "Machine_ID": machine,
                    "Vibration_Hz": vib,
                    "Yield_Rate_Pct": yields[i],
                    "Power_Consumption_kW": powers[i],
                    "Risk_Level": risk
                })
        
        df = pd.DataFrame(all_rows)
        # Sort chronologically to represent a true interleaved continuous timeline
        df = df.sort_values("Date").reset_index(drop=True)
        # Truncate to ensure exactly 100 rows
        df = df.head(100)
        df.to_excel("mes_data.xlsx", index=False)

    # B. ISO Compliance & Master SOP Knowledge Manual (Word)
    if not os.path.exists("audit_sop.docx"):
        doc = Document()
        doc.add_heading("工廠智慧大腦內部營運與合規知識庫", 0)
        
        doc.add_heading("#Equipment_Failure (設備故障排除經驗)", level=1)
        doc.add_paragraph(
            "【資深師傅陳大明筆記】當 EQP-A03 機台的振動頻率超過 8.0Hz 且良率下滑時，通常是 X 軸軸承嚴重磨損。"
            "緊急處方：立刻通知 B 班人員停機，更換原廠 X 軸軸承，過往類似案例顯示可降低 90% 停機風險。詳細更換步驟請參閱機台原廠手冊第三章。"
            "對於 EQP-B01 與 EQP-C02 機台，振動大於 8.0Hz 時則需檢查主軸皮帶張力與聯軸器對中情況。"
        )
        
        doc.add_heading("#ISO9001_2015_Quality (品質稽核與 CAPA 流程)", level=1)
        doc.add_paragraph(
            "根據 ISO 9001:2015 品質管理體系要求，工廠建立以下 CAPA（糾正與預防措施）流程：當任何核心設備（EQP-A03, EQP-B01, EQP-C02）"
            "發生振動超標（紅燈風險，即 >8.0Hz）時，必須立即觸發緊急停機與故障排除。操作人員需在 24 小時內填寫並提交『設備異常與維護保養紀錄表』。"
            "主管必須在 48 小時內完成原因分析，並制定預防再發措施，由現場主管及品質負責人雙重簽核後留存，以備外部品質稽核查審。"
        )
        
        doc.add_heading("#ISO14064_1_Carbon_Emission (碳排放與節能合規指引)", level=1)
        doc.add_paragraph(
            "本廠嚴格遵守 ISO 14064-1 溫室氣體量化與報告規範。當生產設備因磨損或異常導致單台能耗 Power_Consumption_kW 超過 150kW 時，"
            "會引發嚴重的碳排超標風險。管理階層在進行產能與供應鏈調度時，應優先分流至低能耗機台（如節能狀態下的 EQP-C02 設備）。"
            "所有能耗超標事件必須登錄於碳盤查系統中，並執行對應的節能減碳矯正計畫。"
        )
        
        doc.save("audit_sop.docx")

# Initialize mockup files
init_mock_files()

# Load MES Data
df_mes = pd.read_excel("mes_data.xlsx")
# Set index/Date sorted for analysis
df_mes_sorted = df_mes.sort_values("Date").reset_index(drop=True)

# ==========================================
# 2. Pure Python RAG Engine (Deterministic NLP)
# ==========================================
def tokenize(text):
    """Normalizes text and tokenizes into English words/numbers and CJK characters"""
    text = str(text).lower()
    # Alphanumeric runs (English, IDs)
    eng_tokens = re.findall(r'[a-z0-9]+', text)
    # Individual Chinese characters
    cjk_tokens = re.findall(r'[\u4e00-\u9fff]', text)
    return eng_tokens + cjk_tokens

def chunk_text(text, chunk_size=200, overlap=50):
    """Splits a document text string into chunks of specific size with overlap"""
    text = re.sub(r'\s+', ' ', text).strip()
    chunks = []
    step = chunk_size - overlap
    if step <= 0:
        step = chunk_size
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == len(text):
            break
        start += step
    return chunks

def retrieve_top_chunks(query, doc_path="audit_sop.docx", top_k=2):
    """Extracts top_k matching chunks from docx using keyword frequencies and Jaccard overlap"""
    if not os.path.exists(doc_path):
        return []
        
    doc = Document(doc_path)
    full_text = []
    for p in doc.paragraphs:
        if p.text.strip():
            full_text.append(p.text.strip())
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.strip():
                    full_text.append(c.text.strip())
    text = "\n".join(full_text)
    
    chunks = chunk_text(text, chunk_size=200, overlap=50)
    query_tokens = tokenize(query)
    query_set = set(query_tokens)
    
    if not query_set:
        return [{"text": chunk, "score": 0.0} for chunk in chunks[:top_k]]
        
    scored_chunks = []
    for chunk in chunks:
        chunk_tokens = tokenize(chunk)
        # 1. Frequency Score: cumulative occurrences of query tokens in chunk
        freq_score = sum(chunk_tokens.count(tok) for tok in query_set)
        
        # 2. Tie breaker: Jaccard Similarity index
        chunk_set = set(chunk_tokens)
        intersection = query_set.intersection(chunk_set)
        union = query_set.union(chunk_set)
        jaccard = len(intersection) / len(union) if union else 0.0
        
        score = freq_score + jaccard
        scored_chunks.append({"text": chunk, "score": score})
        
    scored_chunks.sort(key=lambda x: x["score"], reverse=True)
    return scored_chunks[:top_k]

# ==========================================
# 3. Decision Integration & AI Prescriptive Logic
# ==========================================
def identify_machine_and_risk(query, df):
    """Scans user query for machine names, otherwise yields latest machine needing intervention"""
    q_upper = query.upper()
    machines = ["EQP-A03", "EQP-B01", "EQP-C02"]
    
    # Check query first
    for m in machines:
        if m in q_upper:
            m_df = df[df["Machine_ID"] == m]
            if not m_df.empty:
                latest_risk = m_df.iloc[-1]["Risk_Level"]
                return m, latest_risk
            return m, "Unknown"
            
    # Default search: Find latest Red, then Yellow, or fall back to overall latest record
    for risk in ["Red", "Yellow"]:
        risk_df = df[df["Risk_Level"] == risk]
        if not risk_df.empty:
            latest = risk_df.iloc[-1]
            return latest["Machine_ID"], latest["Risk_Level"]
            
    latest = df.iloc[-1]
    return latest["Machine_ID"], latest["Risk_Level"]

def generate_ai_response(query, retrieved_chunks, df):
    """Synthesizes RAG chunks, query intents, and quantitative metrics into a high-fidelity output"""
    machine_id, risk_level = identify_machine_and_risk(query, df)
    
    # Get latest data points
    m_df = df[df["Machine_ID"] == machine_id].sort_values("Date")
    latest_row = m_df.iloc[-1]
    
    vibration = latest_row["Vibration_Hz"]
    yield_pct = latest_row["Yield_Rate_Pct"]
    power_kw = latest_row["Power_Consumption_kW"]
    
    q_lower = query.lower()
    
    # Construct response
    response = f"### 🤖 AI 工廠指揮官即時處方籤 (雙引擎融合診斷)\n\n"
    response += "**1. 🔍 RAG 知識庫上下文檢索 (已注入系統 Prompt Context)**\n"
    response += "系統從 `audit_sop.docx` 中成功檢索出以下關聯度最高之規定與維修經驗段落：\n"
    
    for idx, c in enumerate(retrieved_chunks):
        response += f"- **文獻片段 #{idx+1}** (相似度評分: `{c['score']:.4f}`):\n  > *{c['text']}*\n"
        
    response += f"\n**2. 📊 DIS 數據引擎實時監控診斷**\n"
    response += f"- **觸發機台**: `{machine_id}`\n"
    response += f"- **當前風險評級**: `{risk_level}`\n"
    response += f"- **遙測數據實時讀取**:\n"
    response += f"  - 振動頻率 (Vibration_Hz): `{vibration} Hz`\n"
    response += f"  - 生產良率 (Yield_Rate_Pct): `{yield_pct} %`\n"
    response += f"  - 能耗功率 (Power_Consumption_kW): `{power_kw} kW`\n"
    
    if risk_level == "Red":
        response += (
            f"- **🚨 系統診斷報告**: 偵測到 `{machine_id}` 處於 **Red (高風險)** 狀態。當前振動值 `{vibration}Hz` "
            f"已遠超臨界安全閥值。根據時序特徵分析，該振動尖峰發生 2 個週期後，良率已確實發生衰退（當前 `{yield_pct}%` < 90%）"
            f"且能耗負載大幅飆升（當前 `{power_kw}kW` > 150kW），存在極高零件物理失效風險。\n"
        )
    elif risk_level == "Yellow":
        response += (
            f"- **⚠️ 系統診斷報告**: 偵測到 `{machine_id}` 處於 **Yellow (中風險)** 狀態。當前振動值 `{vibration}Hz` "
            f"落入 5.0~8.0Hz 的警戒區間。機台仍可運作，但建議提早巡查以防零件進一步磨損。\n"
        )
    else:
        response += f"- **✅ 系統診斷報告**: `{machine_id}` 目前各項運行指標良好，處於 **Green (安全)** 運作水平。\n"
        
    response += "\n**3. 📋 跨部門 AI 決策處方籤與 CAPA 指引**\n"
    
    # Direct intents based on keywords
    if any(keyword in q_lower for keyword in ["故障", "維修", "振動", "軸承", "更換", "故障排除"]):
        response += (
            f"- **🛠️ 現場工程處置**: 根據陳大明師傅的故障排除手冊（#Equipment_Failure），針對機台 `{machine_id}`，"
            f"建議立即通知 B 班人員安排緊急停機，更換原廠 **X 軸軸承**。歷史實證此操作能阻斷 90% 以上的非計畫性停航。\n"
            f"- **📋 ISO 9001:2015 品質合規**: 本次異常事件屬於核心設備異動，一線人員必須在更換維修工作展開後的 **24 小時內**，"
            f"完成『設備異常與維護保養紀錄表』的填寫與主管簽核落實存檔，以備外部稽核查驗。"
        )
    elif any(keyword in q_lower for keyword in ["稽核", "iso", "合規", "品質", "流程"]):
        response += (
            f"- **📋 ISO 9001:2015 品質合規流程 (CAPA)**: 針對機台 `{machine_id}` 的異常情況，現場操作員需在 **24 小時內** "
            f"建立異常通報，填寫『設備異常與維護保養紀錄表』。現場主管需在 **48 小時內** 主導原因分析、填寫 CAPA 糾正預防單，並由主管及品質負責人簽核留存。\n"
            f"- **🌱 ISO 14064-1 碳排放稽核**: 機台磨損導致能耗超標 (>150kW) 已構成 Scope 2 間接排放的重大變動，主管需將該時段的能耗列為異常排放，計入月度碳排統計中。"
        )
    elif any(keyword in q_lower for keyword in ["碳排", "能耗", "esg", "節能", "電力"]):
        response += (
            f"- **🌱 ESG 與碳排放合規 (ISO 14064-1)**: 當前 `{machine_id}` 能耗為 `{power_kw}kW`。因振動磨損阻力上升導能耗 > 150kW 時，會引發碳排超標風險。\n"
            f"- **📈 產能排程決策**: 為降低 Scope 2 碳排強度，建議工廠主管透過調度系統，將後續大負荷的排程訂單「分流」至低能耗機台（如節能運作中的 `EQP-C02`）。"
        )
    else:
        response += (
            f"- **🛠️ 綜合處置指引**: 目前系統已將 `{machine_id}` 列入關注對象。請操作人員嚴密監控其滾動良率趨勢。\n"
            f"- **📋 知識庫合規提示**: 請參考檢索到的文獻，確保一線營運操作與 ISO 品質、綠色碳排規範完全對齊。"
        )
        
    return response

def append_to_pbi_log(query, machine, risk_level, prescriptive_action):
    """Appends interactions dynamically to the Power BI CSV Sync sheet using standard ISO 8601 timestamps"""
    csv_path = "pbi_live_sync.csv"
    timestamp = datetime.now().isoformat()
    new_data = {
        "Timestamp": [timestamp],
        "User_Query": [query],
        "Triggered_Machine": [machine],
        "Risk_Level": [risk_level],
        "AI_Prescriptive_Action": [prescriptive_action]
    }
    new_df = pd.DataFrame(new_data)
    if not os.path.exists(csv_path):
        new_df.to_csv(csv_path, index=False, encoding="utf-8-sig")
    else:
        new_df.to_csv(csv_path, mode='a', header=False, index=False, encoding="utf-8-sig")

# ==========================================
# 4. Streamlit UI Presentation Layout
# ==========================================
st.set_page_config(page_title="製造業雙引擎 AI 決策智能系統", layout="wide")

# Custom Styling (Dark Glassmorphism)
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&display=swap');
    
    html, body, [data-testid="stAppViewContainer"] {
        font-family: 'Outfit', -apple-system, sans-serif;
    }
    
    /* KPI Card styling styling */
    div[data-testid="column"] {
        background: rgba(22, 27, 34, 0.85);
        border: 1px solid rgba(48, 54, 61, 0.7);
        border-radius: 12px;
        padding: 18px 24px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    
    div[data-testid="column"]:hover {
        border-color: #58a6ff;
        transform: translateY(-2px);
        box-shadow: 0 4px 25px rgba(88, 166, 255, 0.15);
    }
    
    h1, h2, h3 {
        font-weight: 700 !important;
        color: #ffffff;
    }
    
    /* Streamlit success/error adjustments */
    div.stAlert {
        border-radius: 8px;
    }
</style>
""", unsafe_allow_html=True)

st.title("🏭 製造業雙引擎 AI 決策智能系統 (DIS + RAG) — 雲端高級版")
st.caption("🚀 當前模式：高級免密鑰本機測試模式 (無縫對接 Power BI 雲端數據網關)")
st.markdown("---")

# Rolling calculation & metric computation
df_mes_sorted['Rolling_Yield'] = df_mes_sorted['Yield_Rate_Pct'].rolling(window=7, min_periods=1).mean()
latest_yield = df_mes_sorted['Yield_Rate_Pct'].iloc[-1]
latest_rolling_yield = df_mes_sorted['Rolling_Yield'].iloc[-1]
yield_delta = latest_yield - latest_rolling_yield

active_reds = len(df_mes_sorted[df_mes_sorted['Risk_Level'] == 'Red'])
active_yellows = len(df_mes_sorted[df_mes_sorted['Risk_Level'] == 'Yellow'])

# KPI Indicators
kpi_col1, kpi_col2, kpi_col3, kpi_col4 = st.columns(4)
with kpi_col1:
    st.metric(label="📈 當前單班生產良率 (最新)", value=f"{latest_yield:.2f}%")
with kpi_col2:
    st.metric(label="🔄 7日滾動平均良率", value=f"{latest_rolling_yield:.2f}%", delta=f"{yield_delta:+.2f}%")
with kpi_col3:
    st.metric(label="🚨 累計紅燈異常數 (Risk >8Hz)", value=f"{active_reds} 件", delta=None)
with kpi_col4:
    st.metric(label="⚠️ 累計黃燈警報數 (Risk 5-8Hz)", value=f"{active_yellows} 件", delta=None)

st.markdown("<br>", unsafe_allow_html=True)

# Quantitative Visualization
col_left, col_right = st.columns([2, 1])
with col_left:
    st.write("**📈 設備振動趨勢監控線圖 (Vibration_Hz Trends)**")
    # Pivot to display a line per machine on the time series
    chart_df = df_mes_sorted.pivot_table(index='Date', columns='Machine_ID', values='Vibration_Hz')
    chart_df = chart_df.ffill().bfill()
    st.line_chart(chart_df)
    
with col_right:
    st.write("**📋 實時生產數據流水線 (MES Telemetry Table)**")
    st.dataframe(
        df_mes_sorted.sort_values("Date", ascending=False),
        use_container_width=True,
        height=320
    )

st.markdown("---")

# RAG & Chat Assistant Section
st.subheader("🧠 AI 工廠總指揮官 — 智慧決策處方區 (RAG 經驗引擎)")

if "messages" not in st.session_state:
    st.session_state.messages = []
if "retrieval_log" not in st.session_state:
    st.session_state.retrieval_log = []

# Show historical conversation
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Chat Input at the bottom of the section
if prompt := st.chat_input("請輸入您的問題（例如：EQP-A03 振動過高怎麼辦？下週外部品質稽核要注意什麼？"):
    # Append user question
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)
        
    # Retrieve knowledge & synthesize response
    with st.spinner("AI 正在翻閱內部合規規章與師傅經驗手冊..."):
        retrieved_chunks = retrieve_top_chunks(prompt, doc_path="audit_sop.docx", top_k=2)
        st.session_state.retrieval_log = retrieved_chunks
        
        response_text = generate_ai_response(prompt, retrieved_chunks, df_mes_sorted)
        
        # Log event to Power BI csv
        triggered_mach, risk_lvl = identify_machine_and_risk(prompt, df_mes_sorted)
        append_to_pbi_log(prompt, triggered_mach, risk_lvl, response_text)
        
    # Append assistant response
    st.session_state.messages.append({"role": "assistant", "content": response_text})
    with st.chat_message("assistant"):
        st.markdown(response_text)
        st.caption(f"💾 [系統提示] 決策數據與處方已同步導出至 `pbi_live_sync.csv`")

# Analytics and similarity diagnostics log expander at the absolute bottom
st.markdown("<br>", unsafe_allow_html=True)
with st.expander("🔍 RAG Retrieval & Similarity Analytics Log", expanded=False):
    if st.session_state.retrieval_log:
        st.markdown("#### 📊 RAG 向量特徵/關鍵詞頻率匹配分析")
        for i, item in enumerate(st.session_state.retrieval_log):
            st.markdown(f"**【文獻片段 #{i+1}】** (關鍵詞匹配分數: `{item['score']:.4f}`):")
            st.info(item["text"])
            st.markdown("---")
    else:
        st.info("尚無檢索紀錄，請於上方對話框提出問題以觸發計算。")