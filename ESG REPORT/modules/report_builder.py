# -*- coding: utf-8 -*-
"""
ESG 永續報告書自動化生成系統 - Word 排版與組裝模組
"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import config

# 設定 Matplotlib 中文字型，避免亂碼
plt.rcParams['font.sans-serif'] = ['Microsoft JhengHei', 'Microsoft YaHei', 'SimHei', 'Arial Unicode MS', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False  # 正常顯示負號

class ESGReportBuilder:
    """
    使用 python-docx 建立高度結構化的長文件，自動整合 Matplotlib 數據圖表與 AI 文本。
    """
    
    def __init__(self):
        pass

    @staticmethod
    def hex_to_rgb(hex_str):
        """將十六進位顏色碼轉換為 docx.shared.RGBColor"""
        hex_str = hex_str.lstrip('#')
        return RGBColor(*(int(hex_str[i:i+2], 16) for i in (0, 2, 4)))

    @staticmethod
    def add_page_number(run):
        """利用 Word XML 語法動態加入頁碼"""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        r = run._r
        r.append(fldChar1)
        r.append(instrText)
        r.append(fldChar2)
        r.append(fldChar3)

    @staticmethod
    def set_run_fonts(run, font_name="Arial", eastasia_font="Microsoft JhengHei"):
        """設定 Run 的英文與中文字型，確保在不同語系 Office 中均能正確顯示"""
        run.font.name = font_name
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), eastasia_font)
        rPr.append(rFonts)

    def generate_emissions_chart(self, emissions_data, temp_path="temp_emissions.png"):
        """繪製範疇一、範疇二排放比例環狀圖"""
        try:
            em = emissions_data.get("emissions_data", {})
            s1_total = em.get("scope_1_direct", {}).get("總計", 0.0)
            s2_total = em.get("scope_2_indirect", {}).get("總計", 0.0)
            
            labels = ['範疇一 直接排放', '範疇二 間接排放']
            sizes = [s1_total, s2_total]
            colors = [config.COLOR_GRI_GREEN, config.COLOR_TECH_BLUE]
            
            # 建立圖表
            fig, ax = plt.subplots(figsize=(6, 4.5), dpi=300)
            wedges, texts, autotexts = ax.pie(
                sizes, 
                labels=labels, 
                autopct='%1.1f%%', 
                startangle=140, 
                colors=colors,
                wedgeprops=dict(width=0.4, edgecolor='w') # 環狀圖
            )
            
            # 設定字型
            for text in texts:
                text.set_fontsize(10)
            for autotext in autotexts:
                autotext.set_fontsize(10)
                autotext.set_weight('bold')
                autotext.set_color('white')
                
            plt.title(f"溫室氣體範疇排放占比圖 ({emissions_data.get('reporting_year', '2025')} 年度)", fontsize=12, pad=20, weight='bold')
            plt.tight_layout()
            plt.savefig(temp_path, bbox_inches='tight')
            plt.close()
            return temp_path
        except Exception as e:
            print(f"⚠️ 繪製環境圖表失敗: {e}")
            return None

    def generate_social_chart(self, social_data, temp_path="temp_social.png"):
        """繪製人資培訓時數長條圖"""
        try:
            sd = social_data.get("social_data", {})
            tm = sd.get("training_metrics", {})
            
            # 排除全體員工，只取子項目繪製長條圖
            labels = []
            values = []
            for name, item in tm.items():
                if "全體" not in name:
                    labels.append(name.replace("平均培訓時數", ""))
                    values.append(item.get("value", 0.0))
            
            if not labels:
                labels = list(tm.keys())
                values = [item.get("value", 0.0) for item in tm.values()]
                
            fig, ax = plt.subplots(figsize=(7, 4.5), dpi=300)
            
            bars = ax.barh(labels, values, color=config.COLOR_TECH_BLUE, height=0.6)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#cccccc')
            ax.spines['bottom'].set_color('#cccccc')
            ax.xaxis.grid(True, linestyle='--', alpha=0.6, color='#dddddd')
            ax.set_axisbelow(True)
            
            # 加入數值標籤
            for bar in bars:
                width = bar.get_width()
                ax.text(
                    width + 0.5, 
                    bar.get_y() + bar.get_height()/2, 
                    f"{width} 小時", 
                    ha='left', 
                    va='center', 
                    fontsize=9, 
                    weight='bold'
                )
                
            plt.title("各維度員工平均培訓時數對比", fontsize=12, pad=20, weight='bold')
            plt.tight_layout()
            plt.savefig(temp_path, bbox_inches='tight')
            plt.close()
            return temp_path
        except Exception as e:
            print(f"⚠️ 繪製人資圖表失敗: {e}")
            return None

    def build_full_report(self, company_name, year, chapters_data):
        """
        初始化 Word 文件，動態建立大綱、圖表與文本組裝，輸出完整的 Word 報告書。
        """
        doc = Document()
        
        # 1. 頁面設定 (邊界皆為 1 英吋)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            # 設定封面不同頁首頁尾
            section.different_first_page_header_footer = True
            
            # 設定頁首
            header = section.header
            header_p = header.paragraphs[0]
            header_p.text = ""
            header_run = header_p.add_run(f"🌱 {company_name} | {year} 年度企業永續報告書草稿")
            self.set_run_fonts(header_run)
            header_run.font.size = Pt(9)
            header_run.font.color.rgb = RGBColor(128, 128, 128)
            header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 設定頁尾
            footer = section.footer
            footer_p = footer.paragraphs[0]
            footer_p.text = ""
            footer_run_text = footer_p.add_run("企業永續報告書初稿 (地端 AI 自動化生成)  |  頁 ")
            self.set_run_fonts(footer_run_text)
            footer_run_text.font.size = Pt(9)
            footer_run_text.font.color.rgb = RGBColor(128, 128, 128)
            
            footer_run_num = footer_p.add_run()
            self.add_page_number(footer_run_num)
            footer_run_num.font.size = Pt(9)
            footer_run_num.font.color.rgb = RGBColor(128, 128, 128)
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 企業顏色定義
        c_green = self.hex_to_rgb(config.COLOR_GRI_GREEN)
        c_blue = self.hex_to_rgb(config.COLOR_TECH_BLUE)
        c_dark = self.hex_to_rgb(config.COLOR_DARK_TEXT)
        
        # 輔助排版格式化函數
        def add_heading_1(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            self.set_run_fonts(run)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = c_green
            return p

        def add_heading_2(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            self.set_run_fonts(run)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = c_blue
            return p

        def add_body_text(text):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_after = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # 分段處理
            paragraphs = text.split("\n")
            for i, para in enumerate(paragraphs):
                if i > 0:
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.3
                    p.paragraph_format.space_after = Pt(8)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(para.strip())
                self.set_run_fonts(run)
                run.font.size = Pt(11)
                run.font.color.rgb = c_dark
            return p

        # ==================== Page 1: 封面 ====================
        cover_p = doc.add_paragraph()
        cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_p.paragraph_format.space_before = Pt(120)
        
        # 標題
        title_run = cover_p.add_run(f"{company_name}\n\n{year} 年度企業永續報告書草稿\n")
        self.set_run_fonts(title_run)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = c_green
        
        # 副標題
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(180)
        sub_run = sub_p.add_run("GRI 永續性報導準則 (GRI Standards) 合規報告書初稿")
        self.set_run_fonts(sub_run)
        sub_run.font.size = Pt(13)
        sub_run.font.italic = True
        sub_run.font.color.rgb = c_blue
        
        # 底部 metadata
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.paragraph_format.line_spacing = 1.5
        meta_run = meta_p.add_run(
            f"報告年度：{year} 年度\n"
            "報導核心：多指標綜合永續報告書\n"
            "製表單位：ESG 企業永續發展小組\n"
            "發布時間：地端 AI 自動生成系統自動輸出\n"
        )
        self.set_run_fonts(meta_run)
        meta_run.font.size = Pt(10.5)
        meta_run.font.color.rgb = RGBColor(90, 90, 90)
        
        doc.add_page_break()

        # 計算動態頁碼與設定
        current_page = 5
        pages = {}
        
        # 按照順序計算各章節的起始頁碼
        orderOfChapters = ["GRI 2", "GRI 201", "GRI 205", "GRI 302", "GRI 305", "GRI 306", "GRI 401", "GRI 404", "GRI 405"]
        for ch in orderOfChapters:
            if ch in chapters_data:
                sub_count = len(chapters_data[ch].get("sub_chapters", {}))
                if sub_count > 0:
                    pages[ch] = current_page
                    if sub_count >= 3:
                        current_page += 2
                    else:
                        current_page += 1
                    
        page_appendix = current_page

        # ==================== Page 2: 目錄 ====================
        add_heading_1("目錄 (Table of Contents)")
        
        toc_lines = [
            "一、 董事長致詞 ................................................................................................ 3",
            "二、 關於本報告書 ............................................................................................. 4"
        ]
        
        chapter_counter = 3
        chapter_titles = {
            "GRI 2": "GRI 2：一般揭露 (General Disclosures)",
            "GRI 201": "GRI 201：經濟績效 (Economic Performance)",
            "GRI 205": "GRI 205：反貪腐 (Anti-corruption)",
            "GRI 302": "GRI 302：能源消耗 (Energy)",
            "GRI 305": "GRI 305：溫室氣體排放量揭露 (Emissions)",
            "GRI 306": "GRI 306：廢棄物與回收 (Waste)",
            "GRI 401": "GRI 401：員工聘用與流動 (Employment)",
            "GRI 404": "GRI 404：培訓與教育 (Training & Education)",
            "GRI 405": "GRI 405：多元與平等機會 (Diversity & Equality)"
        }
        
        num_map = ["三", "四", "五", "六", "七", "八", "九", "十", "十一", "十二"]
        
        for ch in orderOfChapters:
            if ch in chapters_data and len(chapters_data[ch].get("sub_chapters", {})) > 0:
                ch_num_str = num_map[chapter_counter - 3]
                p_num = pages[ch]
                toc_lines.append(f"{ch_num_str}、 {chapter_titles[ch]} ............................................................................ {p_num}")
                chapter_counter += 1
                
        toc_lines.append(f"附錄、 GRI 揭露指標索引對照表 .............................................................................. {page_appendix}")
        toc_text = "\n".join(toc_lines) + "\n"

        toc_p = doc.add_paragraph()
        toc_p.paragraph_format.line_spacing = 1.6
        toc_run = toc_p.add_run(toc_text)
        self.set_run_fonts(toc_run)
        toc_run.font.size = Pt(11)
        toc_run.font.color.rgb = c_dark
        
        doc.add_page_break()

        # ==================== Page 3: 董事長致詞 ====================
        add_heading_1("一、 董事長致詞")
        
        message_text = (
            "各位夥伴與關心本公司的持份者：\n\n"
            f"在過去的一年中，全球面臨氣候變遷帶來的嚴峻挑戰，ESG 永續發展已不僅是企業的社會責任，更攸關企業在市場中的生存韌性。 "
            f"在 {year} 年度，本公司持續優化內部治理體系，以綠色製造與人本發展為核心雙軸。 "
            f"本報告書詳實揭露了我們在溫室氣體減排 (GRI 305) 以及員工職涯培訓與教育 (GRI 404) 的具體實踐成果。 "
            f"我們深信，透過落實低碳製程與優質的人才教育計畫，我們將引領企業航向永續淨零的綠色未來。"
        )
        add_body_text(message_text)
        
        doc.add_page_break()

        # ==================== Page 4: 關於本報告書 ====================
        add_heading_1("二、 關於本報告書")
        
        about_text = (
            f"本報告書為 {company_name} 發布之年度企業永續報告書，揭露期間為西元 {year} 年 1 月 1 日至 12 月 31 日止之永續績效。 "
            "本報告書依循全球永續性標準理事會 (GSSB) 發布之 GRI 永續報導準則 (GRI Standards) 進行編撰，重點揭露環境面向之排放 (GRI 305) 及社會面向之培訓與教育 (GRI 404) 指標。 "
            "本報告書所包含之財務與非財務數據，皆經由內部跨部門小組進行收集、清洗與確信審核，並以地端 AI 安全防護生成核心質化內文。 "
            "未來本報告書將維持每年定期發布，以向社會大眾展示本公司實踐永續發展之堅定承諾。"
        )
        add_body_text(about_text)
        
        doc.add_page_break()

        # ==================== GRI 2 ====================
        if "GRI 2" in chapters_data:
            data = chapters_data["GRI 2"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("三、 GRI 2：一般揭露與企業永續治理概況")
                
                if "2.2.1" in sub_chapters:
                    add_heading_2("2.2.1 組織基本概況、據點與資本規模")
                    add_body_text(sub_chapters["2.2.1"])
                if "2.2.2" in sub_chapters:
                    add_heading_2("2.2.2 商業活動與供應鏈價值關係")
                    add_body_text(sub_chapters["2.2.2"])
                if "2.2.3" in sub_chapters:
                    add_heading_2("2.2.3 員工聘用特性與人力資源分佈")
                    add_body_text(sub_chapters["2.2.3"])
                
                add_heading_2("2.2.4 組織基本概況數據表")
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '組織特性項目'
                hdr_cells[1].text = '揭露內容說明'
                for cell in hdr_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            self.set_run_fonts(r)
                            r.font.bold = True
                            r.font.size = Pt(10)
                
                gd_items = data.get("general_data", {})
                for k, v in gd_items.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(k).replace("_", " ")
                    row_cells[1].text = str(v)
                
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                self.set_run_fonts(r)
                                r.font.size = Pt(9.5)
                
                doc.add_page_break()

        # ==================== GRI 201 ====================
        if "GRI 201" in chapters_data:
            data = chapters_data["GRI 201"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("四、 GRI 201：經濟績效")
                
                if "2.1.1" in sub_chapters:
                    add_heading_2("2.1.1 直接產生與分配之經濟價值分析")
                    add_body_text(sub_chapters["2.1.1"])
                if "2.1.2" in sub_chapters:
                    add_heading_2("2.1.2 氣候變遷對企業營運之財務衝擊")
                    add_body_text(sub_chapters["2.1.2"])
                
                add_heading_2("2.1.3 直接產生與分配的經濟價值表")
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '經濟價值組成項目'
                hdr_cells[1].text = '金額 (萬元)'
                for cell in hdr_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            self.set_run_fonts(r)
                            r.font.bold = True
                            r.font.size = Pt(10)
                
                ec_items = data.get("economic_data", {})
                for k, v in ec_items.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(k).replace("_", " ")
                    row_cells[1].text = f"{v:,.2f}" if isinstance(v, (int, float)) else str(v)
                
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                self.set_run_fonts(r)
                                r.font.size = Pt(9.5)
                
                doc.add_page_break()

        # ==================== GRI 205 ====================
        if "GRI 205" in chapters_data:
            data = chapters_data["GRI 205"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("五、 GRI 205：反貪腐")
                
                if "2.5.1" in sub_chapters:
                    add_heading_2("2.5.1 反貪腐政策傳達、簽署與培訓統計")
                    add_body_text(sub_chapters["2.5.1"])
                if "2.5.2" in sub_chapters:
                    add_heading_2("2.5.2 誠信經營確立事件與檢舉防範機制")
                    add_body_text(sub_chapters["2.5.2"])
                
                add_heading_2("2.5.3 反貪腐與誠信經營運作指標表")
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '誠信經營防範指標'
                hdr_cells[1].text = '統計數值'
                hdr_cells[2].text = '單位'
                for cell in hdr_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            self.set_run_fonts(r)
                            r.font.bold = True
                            r.font.size = Pt(10)
                
                ac_items = data.get("anti_corruption_data", {})
                for k, v in ac_items.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(k).split("_")[0]
                    row_cells[1].text = str(v)
                    row_cells[2].text = str(k).split("_")[-1] if "_" in str(k) else ""
                
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                self.set_run_fonts(r)
                                r.font.size = Pt(9.5)
                
                doc.add_page_break()

        # ==================== GRI 302 ====================
        if "GRI 302" in chapters_data:
            data = chapters_data["GRI 302"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("六、 GRI 302：能源消耗")
                
                if "3.2.1" in sub_chapters:
                    add_heading_2("3.2.1 組織內部能源消耗數據解讀")
                    add_body_text(sub_chapters["3.2.1"])
                if "3.2.2" in sub_chapters:
                    add_heading_2("3.2.2 能源密集度與節能減量成效")
                    add_body_text(sub_chapters["3.2.2"])
                
                add_heading_2("3.2.3 企業能源消耗統計表")
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '能源種類'
                hdr_cells[1].text = '消耗數據'
                hdr_cells[2].text = '單位'
                for cell in hdr_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            self.set_run_fonts(r)
                            r.font.bold = True
                            r.font.size = Pt(10)
                
                en_items = data.get("energy_data", {})
                for k, v in en_items.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(k).split("_")[0]
                    row_cells[1].text = f"{v:,.1f}" if isinstance(v, (int, float)) else str(v)
                    row_cells[2].text = str(k).split("_")[-1] if "_" in str(k) else ""
                
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                self.set_run_fonts(r)
                                r.font.size = Pt(9.5)
                
                doc.add_page_break()

        # ==================== GRI 305 ====================
        if "GRI 305" in chapters_data:
            data = chapters_data["GRI 305"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("七、 GRI 305：溫室氣體排放量揭露")
                
                if "3.1.1" in sub_chapters:
                    add_heading_2("3.1.1 範疇一（直接溫室氣體排放）來源與數據深度解讀")
                    add_body_text(sub_chapters["3.1.1"])
                    
                if "3.1.2" in sub_chapters:
                    add_heading_2("3.1.2 範疇二（能源間接溫室氣體排放）外購電力分析與減碳路徑")
                    add_body_text(sub_chapters["3.1.2"])
                
                if "3.1.1" in sub_chapters or "3.1.2" in sub_chapters:
                    add_heading_2("3.1.2.2 溫室氣體排放細部數據表")
                    
                    em_data = data.get("emissions_data", {})
                    s1 = em_data.get("scope_1_direct", {})
                    s2 = em_data.get("scope_2_indirect", {})
                    
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Light Shading Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '範疇別'
                    hdr_cells[1].text = '排放源名稱'
                    hdr_cells[2].text = '碳排放量 (公噸 CO2e)'
                    hdr_cells[3].text = '百分比 (%)'
                    
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                self.set_run_fonts(run)
                                run.font.bold = True
                                run.font.size = Pt(10)
                    
                    total_val = em_data.get("total_emissions_tCO2e", 1.0)
                    
                    for name, val in s1.items():
                        if name == "總計":
                            continue
                        row_cells = table.add_row().cells
                        row_cells[0].text = "範疇一 (直接)"
                        row_cells[1].text = name
                        row_cells[2].text = f"{val:,}"
                        pct = round((val / total_val) * 100, 1) if total_val > 0 else 0.0
                        row_cells[3].text = f"{pct}%"
                        
                    for name, val in s2.items():
                        if name == "總計":
                            continue
                        row_cells = table.add_row().cells
                        row_cells[0].text = "範疇二 (間接)"
                        row_cells[1].text = name
                        row_cells[2].text = f"{val:,}"
                        pct = round((val / total_val) * 100, 1) if total_val > 0 else 0.0
                        row_cells[3].text = f"{pct}%"
                        
                    row_cells = table.add_row().cells
                    row_cells[0].text = "範疇一與範疇二總計"
                    row_cells[1].text = "所有排放源"
                    row_cells[2].text = f"{total_val:,}"
                    row_cells[3].text = "100.0%"
                    
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    self.set_run_fonts(run)
                                    run.font.size = Pt(9.5)
                    
                    p_space = doc.add_paragraph()
                    p_space.paragraph_format.space_before = Pt(10)
                    
                    doc.add_page_break()
                
                if "3.1.3" in sub_chapters:
                    add_heading_2("3.1.3 年度排放變動率（YoY）與減量成效評估")
                    add_body_text(sub_chapters["3.1.3"])
                    
                    add_heading_2("3.1.3.2 排放源與範疇占比可視化")
                    
                    chart_path = self.generate_emissions_chart(data)
                    if chart_path and os.path.exists(chart_path):
                        img_p = doc.add_paragraph()
                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_run = img_p.add_run()
                        img_run.add_picture(chart_path, width=Inches(5.0))
                        
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.paragraph_format.space_after = Pt(20)
                        cap_run = caption_p.add_run(f"圖 3-1 {year} 年度環境範疇一與範疇二排放分布 donut 圖")
                        self.set_run_fonts(cap_run)
                        cap_run.font.size = Pt(9)
                        cap_run.font.italic = True
                        
                        os.remove(chart_path)
                    
                    doc.add_page_break()

        # ==================== GRI 306 ====================
        if "GRI 306" in chapters_data:
            data = chapters_data["GRI 306"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("八、 GRI 306：廢棄物與回收")
                
                if "3.6.1" in sub_chapters:
                    add_heading_2("3.6.1 廢棄物產生源與源頭減量措施")
                    add_body_text(sub_chapters["3.6.1"])
                if "3.6.2" in sub_chapters:
                    add_heading_2("3.6.2 廢棄物回收與循環再利用效益")
                    add_body_text(sub_chapters["3.6.2"])
                if "3.6.3" in sub_chapters:
                    add_heading_2("3.6.3 廢棄物最終處置與合規評估")
                    add_body_text(sub_chapters["3.6.3"])
                
                add_heading_2("3.6.4 廢棄物產生與處置方式明細表")
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '廢棄物與處置類別'
                hdr_cells[1].text = '統計值'
                hdr_cells[2].text = '單位'
                for cell in hdr_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            self.set_run_fonts(r)
                            r.font.bold = True
                            r.font.size = Pt(10)
                
                wt_items = data.get("waste_data", {})
                for k, v in wt_items.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(k).split("_")[0]
                    row_cells[1].text = str(v)
                    row_cells[2].text = str(k).split("_")[-1] if "_" in str(k) else "噸"
                
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                self.set_run_fonts(r)
                                r.font.size = Pt(9.5)
                
                doc.add_page_break()

        # ==================== GRI 401 ====================
        if "GRI 401" in chapters_data:
            data = chapters_data["GRI 401"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("九、 GRI 401：員工聘用與流動")
                
                if "4.1.1.b" in sub_chapters:
                    add_heading_2("4.1.1.b 員工流動率與新進率結構解讀")
                    add_body_text(sub_chapters["4.1.1.b"])
                if "4.1.2" in sub_chapters:
                    add_heading_2("4.1.2 關懷福利政策與育嬰留停成效")
                    add_body_text(sub_chapters["4.1.2"])
                
                add_heading_2("4.1.3 員工聘用與流動指標統計表")
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '人力資源流動指標'
                hdr_cells[1].text = '統計值'
                hdr_cells[2].text = '單位'
                for cell in hdr_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            self.set_run_fonts(r)
                            r.font.bold = True
                            r.font.size = Pt(10)
                
                em_items = data.get("employment_data", {})
                for k, v in em_items.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(k).split("_")[0]
                    row_cells[1].text = str(v)
                    row_cells[2].text = str(k).split("_")[-1] if "_" in str(k) else "人"
                
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                self.set_run_fonts(r)
                                r.font.size = Pt(9.5)
                
                doc.add_page_break()

        # ==================== GRI 404 ====================
        if "GRI 404" in chapters_data:
            data = chapters_data["GRI 404"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("十、 GRI 404：培訓與教育")
                
                if "4.1.1" in sub_chapters:
                    add_heading_2("4.1.1 員工平均培訓時數結構分析（依職級與性別）")
                    add_body_text(sub_chapters["4.1.1"])
                    
                if "4.1.2" in sub_chapters:
                    add_heading_2("4.1.2 員工技能提升與過渡協助計畫執行效益")
                    add_body_text(sub_chapters["4.1.2"])
                
                if "4.1.1" in sub_chapters or "4.1.2" in sub_chapters:
                    add_heading_2("4.1.2.2 培訓與教育績效統計表")
                    
                    sd = data.get("social_data", {})
                    tm = sd.get("training_metrics", {})
                    to = sd.get("turnover_metrics", {})
                    
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Light Shading Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '數據指標分類'
                    hdr_cells[1].text = '細部指標名稱'
                    hdr_cells[2].text = '統計值'
                    hdr_cells[3].text = '計量單位'
                    
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                self.set_run_fonts(run)
                                run.font.bold = True
                                run.font.size = Pt(10)
                                
                    for name, val_dict in tm.items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = "培育成長 (Training)"
                        row_cells[1].text = name
                        row_cells[2].text = str(val_dict.get("value"))
                        row_cells[3].text = val_dict.get("unit")
                        
                    for name, val_dict in to.items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = "組織穩定 (Turnover)"
                        row_cells[1].text = name
                        row_cells[2].text = str(val_dict.get("value"))
                        row_cells[3].text = val_dict.get("unit")
                        
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    self.set_run_fonts(run)
                                    run.font.size = Pt(9.5)
                                    
                    p_space = doc.add_paragraph()
                    p_space.paragraph_format.space_before = Pt(10)
                    
                    doc.add_page_break()
                
                if "4.1.3" in sub_chapters:
                    add_heading_2("4.1.3 組織人才穩定度與流動率指標解讀")
                    add_body_text(sub_chapters["4.1.3"])
                    
                    add_heading_2("4.1.3.2 培訓效益與流動性分析")
                    
                    chart_path = self.generate_social_chart(data)
                    if chart_path and os.path.exists(chart_path):
                        img_p = doc.add_paragraph()
                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_run = img_p.add_run()
                        img_run.add_picture(chart_path, width=Inches(5.2))
                        
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.paragraph_format.space_after = Pt(20)
                        cap_run = caption_p.add_run(f"圖 4-1 {year} 年度各維度平均培訓時數對比長條圖")
                        self.set_run_fonts(cap_run)
                        cap_run.font.size = Pt(9)
                        cap_run.font.italic = True
                        
                        os.remove(chart_path)
                    
                    doc.add_page_break()

        # ==================== GRI 405 ====================
        if "GRI 405" in chapters_data:
            data = chapters_data["GRI 405"]
            sub_chapters = data.get("sub_chapters", {})
            if sub_chapters:
                add_heading_1("十一、 GRI 405：多元與平等機會")
                
                if "4.5.1" in sub_chapters:
                    add_heading_2("4.5.1 治理機構與員工結構多元化比例")
                    add_body_text(sub_chapters["4.5.1"])
                if "4.5.2" in sub_chapters:
                    add_heading_2("4.5.2 男女同工同酬與平等晉升機會")
                    add_body_text(sub_chapters["4.5.2"])
                
                add_heading_2("4.5.3 管理與基層員工多元化結構表")
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '多元化組成指標'
                hdr_cells[1].text = '佔比'
                hdr_cells[2].text = '單位'
                for cell in hdr_cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            self.set_run_fonts(r)
                            r.font.bold = True
                            r.font.size = Pt(10)
                
                dv_items = data.get("diversity_data", {})
                for k, v in dv_items.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(k).split("_")[0]
                    row_cells[1].text = str(v)
                    row_cells[2].text = str(k).split("_")[-1] if "_" in str(k) else "%"
                
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                self.set_run_fonts(r)
                                r.font.size = Pt(9.5)
                
                doc.add_page_break()

        # ==================== 附錄 ====================
        add_heading_1("附錄、 GRI 揭露指標索引對照表")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'GRI 準則項目'
        hdr_cells[1].text = '揭露指標項目'
        hdr_cells[2].text = '對應報告章節'
        hdr_cells[3].text = '頁碼標示'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self.set_run_fonts(run)
                    run.font.bold = True
                    run.font.size = Pt(10)
                    
        indexes = []
        
        if "GRI 2" in chapters_data:
            sub_2 = chapters_data["GRI 2"].get("sub_chapters", {})
            if "2.2.1" in sub_2:
                indexes.append(("GRI 2: General Disclosures", "2-1 組織基本概況與據點", "GRI 2 - 2.2.1 組織基本概況與資本規模", str(pages.get("GRI 2", ""))))
            if "2.2.2" in sub_2:
                indexes.append(("GRI 2: General Disclosures", "2-6 商業活動與供應鏈價值關係", "GRI 2 - 2.2.2 商業活動與供應鏈", str(pages.get("GRI 2", ""))))
            if "2.2.3" in sub_2:
                indexes.append(("GRI 2: General Disclosures", "2-7 員工聘用特性與人力資源分佈", "GRI 2 - 2.2.3 員工聘用特性與人力資源", str(pages.get("GRI 2", ""))))
                
        if "GRI 201" in chapters_data:
            sub_201 = chapters_data["GRI 201"].get("sub_chapters", {})
            if "2.1.1" in sub_201:
                indexes.append(("GRI 201: Economic Performance", "201-1 直接產生與分配的經濟價值", "GRI 201 - 2.1.1 直接產生與分配經濟價值", str(pages.get("GRI 201", ""))))
            if "2.1.2" in sub_201:
                indexes.append(("GRI 201: Economic Performance", "201-2 氣候變遷對企業營運之財務衝擊", "GRI 201 - 2.1.2 氣候變遷財務衝擊", str(pages.get("GRI 201", ""))))
                
        if "GRI 205" in chapters_data:
            sub_205 = chapters_data["GRI 205"].get("sub_chapters", {})
            if "2.5.1" in sub_205:
                indexes.append(("GRI 205: Anti-corruption", "205-2 反貪腐政策傳達、簽署與培訓", "GRI 205 - 2.5.1 反貪腐政策與培訓", str(pages.get("GRI 205", ""))))
            if "2.5.2" in sub_205:
                indexes.append(("GRI 205: Anti-corruption", "205-3 誠信經營確立事件與檢舉防範", "GRI 205 - 2.5.2 誠信經營與檢舉機制", str(pages.get("GRI 205", ""))))
                
        if "GRI 302" in chapters_data:
            sub_302 = chapters_data["GRI 302"].get("sub_chapters", {})
            if "3.2.1" in sub_302:
                indexes.append(("GRI 302: Energy", "302-1 組織內部的能源消耗量", "GRI 302 - 3.2.1 組織內部能源消耗數據", str(pages.get("GRI 302", ""))))
            if "3.2.2" in sub_302:
                indexes.append(("GRI 302: Energy", "302-3/4 能源密集度與節能減量成效", "GRI 302 - 3.2.2 能源密集度與節能減量", str(pages.get("GRI 302", ""))))
                
        if "GRI 305" in chapters_data:
            sub_305 = chapters_data["GRI 305"].get("sub_chapters", {})
            if "3.1.1" in sub_305:
                indexes.append(("GRI 305: Emissions 2016", "305-1 直接（範疇一）溫室氣體排放", "GRI 305 - 3.1.1 範疇一來源深度解讀", str(pages.get("GRI 305", ""))))
            if "3.1.2" in sub_305:
                indexes.append(("GRI 305: Emissions 2016", "305-2 能源間接（範疇二）溫室氣體排放", "GRI 305 - 3.1.2 範疇二電力與減碳路徑", str(pages.get("GRI 305", ""))))
            if "3.1.3" in sub_305:
                page_offset = 1 if ("3.1.1" in sub_305 or "3.1.2" in sub_305) else 0
                indexes.append(("GRI 305: Emissions 2016", "305-4 溫室氣體排放密集度/YoY成效", "GRI 305 - 3.1.3 YoY排放變動評估", str(pages.get("GRI 305", 0) + page_offset)))
                
        if "GRI 306" in chapters_data:
            sub_306 = chapters_data["GRI 306"].get("sub_chapters", {})
            if "3.6.1" in sub_306:
                indexes.append(("GRI 306: Waste", "306-3 廢棄物產生源與源頭減量", "GRI 306 - 3.6.1 廢棄物產生源與源頭減量", str(pages.get("GRI 306", ""))))
            if "3.6.2" in sub_306:
                indexes.append(("GRI 306: Waste", "306-4 廢棄物回收與循環再利用", "GRI 306 - 3.6.2 廢棄物回收與循環再利用", str(pages.get("GRI 306", ""))))
            if "3.6.3" in sub_306:
                indexes.append(("GRI 306: Waste", "306-5 廢棄物最終處置與合規評估", "GRI 306 - 3.6.3 最終處置與合規評估", str(pages.get("GRI 306", ""))))
                
        if "GRI 401" in chapters_data:
            sub_401 = chapters_data["GRI 401"].get("sub_chapters", {})
            if "4.1.1.b" in sub_401:
                indexes.append(("GRI 401: Employment", "401-1 員工流動率與新進率結構解讀", "GRI 401 - 4.1.1.b 員工流動與新進率", str(pages.get("GRI 401", ""))))
            if "4.1.2" in sub_401:
                indexes.append(("GRI 401: Employment", "401-2/3 關懷福利政策與育嬰留停成效", "GRI 401 - 4.1.2 關懷福利政策與育嬰留停", str(pages.get("GRI 401", ""))))
                
        if "GRI 404" in chapters_data:
            sub_404 = chapters_data["GRI 404"].get("sub_chapters", {})
            if "4.1.1" in sub_404:
                indexes.append(("GRI 404: Training & Education", "404-1 每名員工每年接受培訓的平均時數", "GRI 404 - 4.1.1 培訓時數結構分析", str(pages.get("GRI 404", ""))))
            if "4.1.2" in sub_404:
                indexes.append(("GRI 404: Training & Education", "404-2 員工技能提升計畫與過渡協助計畫", "GRI 404 - 4.1.2 技能提升與過渡計畫", str(pages.get("GRI 404", ""))))
            if "4.1.3" in sub_404:
                page_offset = 1 if ("4.1.1" in sub_404 or "4.1.2" in sub_404) else 0
                indexes.append(("GRI 404: Training & Education", "404-3 接受定期績效及職涯發展檢核", "GRI 404 - 4.1.3 組織人才穩定度", str(pages.get("GRI 404", 0) + page_offset)))
                
        if "GRI 405" in chapters_data:
            sub_405 = chapters_data["GRI 405"].get("sub_chapters", {})
            if "4.5.1" in sub_405:
                indexes.append(("GRI 405: Diversity & Equal Opportunity", "405-1 治理機構與員工的多元化比例", "GRI 405 - 4.5.1 治理機構與員工多元化", str(pages.get("GRI 405", ""))))
            if "4.5.2" in sub_405:
                indexes.append(("GRI 405: Diversity & Equal Opportunity", "405-2 男女同工同酬與平等晉升機會", "GRI 405 - 4.5.2 男女同工同酬與晉升", str(pages.get("GRI 405", ""))))
        
        for item in indexes:
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
            row_cells[3].text = item[3]
            
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self.set_run_fonts(run)
                        run.font.size = Pt(9)
                        
        # 儲存報告書
        out_filename = f"{company_name}_ESG_Report_{year}.docx"
        out_path = os.path.join(config.OUTPUT_DIR, out_filename)
        doc.save(out_path)
        return out_path
