"""
main.py — 企業數位韌性核心後端 API
功能：接收異常數據 → 去識別化 → RAG 檢索 → 本地 AI 推理 → 生成 CAPA 報告
"""

import os
import re
import uuid
import logging
import asyncio
from datetime import datetime, timedelta
from typing import Optional
from dotenv import load_dotenv

# 載入環境變數
_MAIN_DIR = os.path.dirname(os.path.abspath(__file__))
_DOTENV_PATH = os.path.join(_MAIN_DIR, "..", ".env")
if os.path.exists(_DOTENV_PATH):
    load_dotenv(_DOTENV_PATH)
else:
    load_dotenv()

# 確保根目錄與 backend 目錄皆在 sys.path 中，使 guardrails 與 scenario_config 能被正確 import
import sys as _sys
_ROOT_DIR = os.path.dirname(_MAIN_DIR)
if _ROOT_DIR not in _sys.path:
    _sys.path.insert(0, _ROOT_DIR)
if _MAIN_DIR not in _sys.path:
    _sys.path.insert(0, _MAIN_DIR)
from scenario_config import SCENARIOS as DEMO_SCENARIOS

import chromadb
import httpx
import requests
import google.genai
from google import genai
from fastapi import FastAPI, BackgroundTasks, HTTPException, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field
from chromadb.utils import embedding_functions

# [FEATURE] 匯出正式報告功能
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.responses import StreamingResponse

# [FEATURE] RBAC 權限控制與 JWT 機制 (直接在 main.py 實作以簡化並統一權限系統)
from fastapi.security import OAuth2PasswordRequestForm, HTTPBearer, HTTPAuthorizationCredentials
from jose import jwt, JWTError

SECRET_KEY = "super_secret_jwt_key_for_compliance_audit_system"
ALGORITHM = "HS256"
security = HTTPBearer()

# 模擬的資料庫字典 USERS，包含三個帳號與對應角色
USERS = {
    "admin": {"password": "admin123", "role": "admin"},
    "qms": {"password": "qms123", "role": "qms_manager"},
    "ciso": {"password": "ciso123", "role": "ciso"}
}

class User(BaseModel):
    username: str
    role: str

class LoginRequest(BaseModel):
    username: str
    password: str

async def get_current_user_role(credentials: HTTPAuthorizationCredentials = Depends(security)) -> str:
    """依賴注入函式：驗證 Token 並回傳 Role"""
    token = credentials.credentials
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        role = payload.get("role")
        if not role:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Token 缺少角色資訊"
            )
        return role
    except JWTError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="無效或已過期的 Token",
            headers={"WWW-Authenticate": "Bearer"}
        )

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> User:
    """依賴注入函式：驗證 Token 並回傳 User (相容原先路由)"""
    token = credentials.credentials
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username = payload.get("sub")
        role = payload.get("role")
        if not username or not role:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Token 缺少必要欄位"
            )
        return User(username=username, role=role)
    except JWTError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="無效或已過期的 Token",
            headers={"WWW-Authenticate": "Bearer"}
        )

# ── 日誌設定 ─────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)
log = logging.getLogger(__name__)

# ── 常數與資料庫相對路徑設定 ──────────────────────────────────
# 動態計算相對於當前工作目錄的相對路徑，避免寫死絕對路徑，確保地端與 Streamlit Cloud 均能讀取
_CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
_TARGET_DB_PATH = os.path.join(_CURRENT_DIR, "..", "audit_db")
if not os.path.exists(_TARGET_DB_PATH):
    # 若上一層目錄沒有，嘗試尋找同目錄 (backend) 下的 audit_db
    _TARGET_DB_PATH = os.path.join(_CURRENT_DIR, "audit_db")

# 計算出絕對路徑，傳遞給 ChromaDB
DB_PATH      = os.path.abspath(_TARGET_DB_PATH)
COLLECTION   = "compliance_rules"
OLLAMA_URL   = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:14b")
EMBED_MODEL  = "nomic-embed-text"

# ── 資料庫與 RAG 知識庫初始化 ──────────────────────────────────────
from database import init_db, save_report, get_report, list_reports, get_reports_for_recovery, get_db_stats
from rag_service import init_mock_data
init_db()
init_mock_data()

# ── FastAPI 初始化 ────────────────────────────────────────────
app = FastAPI(
    title="企業數位韌性核心後端",
    description="地端私有化 AI 合規稽核系統 | ISO 14064-1 × ISO 27001 × 生管 SOP",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── 全域報告快取（生產環境請改用 Redis 或資料庫） ────────────
report_store: dict[str, dict] = {}

# ── 全域非同步訊號量 (限制 Ollama 推理 Concurrency = 1) ───────────
ai_semaphore = asyncio.Semaphore(1)


# ╔══════════════════════════════════════════════════════════╗
# ║  資料模型（Pydantic Schemas）                            ║
# ╚══════════════════════════════════════════════════════════╝

# [DEMO FEATURE] 放寬碳強度上限至 100 以相容情境三的 24.75 kgCO2e 數據
class AuditPayload(BaseModel):
    carbon_intensity:   float   = Field(..., ge=0, le=100, description="碳排放強度 kgCO₂e/unit")
    user_id:            str     = Field(..., min_length=1, description="操作人員 ID")
    event_log:          str     = Field(..., description="資安事件日誌原始內容")
    department:         Optional[str] = Field(None, description="部門代碼")
    shift:              Optional[str] = Field(None, description="班別（A/B/C）")
    equipment_id:       Optional[str] = Field(None, description="設備編號")
    production_yield:   Optional[float] = Field(None, ge=0, le=100, description="良率 %")

class AuditStatus(BaseModel):
    task_id:    str
    status:     str
    created_at: str
    report:     Optional[dict] = None

# [FEATURE] HITL 審核機制
class ApprovePayload(BaseModel):
    approver: str = Field(..., description="審核人姓名/員工ID")

# [FEATURE] HITL 審核機制
class RejectPayload(BaseModel):
    rejected_by: str = Field(..., description="退回人姓名/員工ID")
    reason: str = Field("退回重擬", description="退回原因")


# ╔══════════════════════════════════════════════════════════╗
# ║  核心功能模組                                            ║
# ╚══════════════════════════════════════════════════════════╝

def anonymize(payload: AuditPayload) -> dict:
    """
    資料去識別化（Data Anonymization）
    ─────────────────────────────────
    規則：
    1. 員工 ID → 保留部門前綴 + 隱碼後半
    2. 設備 ID → 保留設備類型 + 序號隱碼
    3. 事件日誌 → 移除 IP 位址、email、身分證字號
    """
    # 1. 員工 ID 匿名化
    uid = payload.user_id
    if len(uid) > 4:
        safe_uid = uid[:2] + "***" + uid[-2:]
    else:
        safe_uid = "ANON_OP"

    # 2. 日誌去識別化：移除 IP、email
    safe_log = payload.event_log
    safe_log = re.sub(r"\b\d{1,3}(\.\d{1,3}){3}\b", "[IP_REMOVED]", safe_log)
    safe_log = re.sub(r"[\w\.-]+@[\w\.-]+\.\w+", "[EMAIL_REMOVED]", safe_log)
    safe_log = re.sub(r"[A-Z]\d{9}", "[ID_REMOVED]", safe_log)

    # 3. 設備 ID 匿名化
    safe_eq = payload.equipment_id
    if safe_eq and len(safe_eq) > 4:
        safe_eq = safe_eq[:3] + "***"

    return {
        "user_id":         safe_uid,
        "event_log":       safe_log,
        "equipment_id":    safe_eq,
        "carbon_intensity": payload.carbon_intensity,
        "department":      payload.department or "UNKNOWN_DEPT",
        "shift":           payload.shift or "UNKNOWN_SHIFT",
        "yield_rate":      payload.production_yield,
    }


def risk_assessment(carbon: float, log_text: str, yield_rate: Optional[float]) -> dict:
    """
    規則式風險評級引擎（Rule-Based Risk Engine）
    在 AI 推理之前先快速評級，確保緊急事件立即觸發
    """
    level   = 1
    reasons = []
    flags   = []

    # 碳排放評估
    if carbon > 2.0:
        level = max(level, 3)
        reasons.append(f"碳排放強度 {carbon:.2f} 嚴重超標（上限 1.0，超標 {carbon-1.0:.2f}）")
        flags.append("CARBON_CRITICAL")
    elif carbon > 1.0:
        level = max(level, 2)
        reasons.append(f"碳排放強度 {carbon:.2f} 超標（上限 1.0）")
        flags.append("CARBON_EXCEED")

    # 資安事件評估
    midnight_keywords = ["midnight", "deep night", "深夜", "23:", "00:", "01:", "02:", "03:", "04:", "05:"]
    download_keywords = ["download", "下載", "export", "匯出"]
    if any(k in log_text.lower() for k in midnight_keywords) and \
       any(k in log_text.lower() for k in download_keywords):
        level = max(level, 3)
        reasons.append("偵測到深夜大量下載行為，符合 ISO 27001 A.8.16 重大資安事件判定條件")
        flags.append("SECURITY_MIDNIGHT_DOWNLOAD")

    # 良率評估
    if yield_rate is not None and yield_rate < 85:
        level = max(level, 2)
        reasons.append(f"生產良率 {yield_rate:.1f}% 低於目標值 85%")
        flags.append("YIELD_BELOW_TARGET")

    level_map = {1: "Level-1 輕微", 2: "Level-2 嚴重", 3: "Level-3 緊急停線"}
    response_time = {1: "4 小時內提交 CAPA", 2: "1 小時內初評 / 24 小時完成根因分析", 3: "立即通報廠長 + 資安稽核"}

    return {
        "risk_level":    level,
        "risk_label":    level_map.get(level, "未知"),
        "flags":         flags,
        "reasons":       reasons,
        "response_time": response_time.get(level, ""),
    }


def retrieve_context(query: str, n_results: int = 5) -> str:
    """
    從 ChromaDB 向量資料庫檢索相關 ISO 條文與 SOP
    """
    try:
        ef = embedding_functions.DefaultEmbeddingFunction()
        client = chromadb.PersistentClient(path=DB_PATH)
        collection = client.get_or_create_collection(
            name=COLLECTION,
            embedding_function=ef
        )
        results = collection.query(query_texts=[query], n_results=n_results)
        docs = results.get("documents", [[]])[0]
        metas = results.get("metadatas", [[]])[0]

        context_parts = []
        for doc, meta in zip(docs, metas):
            source = meta.get("source", "unknown")
            context_parts.append(f"【來源：{source}】\n{doc}")

        return "\n\n---\n\n".join(context_parts)

    except Exception as e:
        log.warning(f"⚠️  向量檢索失敗，使用備援知識：{e}")
        return (
            "ISO 14064-1：碳排放強度上限為 1.0 kgCO₂e/unit，超標觸發 CAPA 流程。\n"
            "ISO 27001 A.8.16：深夜異常大量下載視為重大資安事件，須立即隔離帳號。\n"
            "SOP-PROD-001：製程異常 4 小時內提交 CAPA 報告。"
        )


async def call_ai_inference(prompt: str) -> str:
    """
    動態容錯雙推論引擎 (Try-Except 雙軌制)
    優先嘗試調用本地的 Ollama (qwen2.5:14b)，並設定 3 秒的短連線超時 (connect timeout)。
    如果發生連線錯誤或超時，自動無痛切換至雲端備援方案 (google-genai SDK 呼叫 gemini-2.5-flash)。
    """
    async with ai_semaphore:
        try:
            log.info(f"🤖 嘗試呼叫本地 Ollama ({OLLAMA_MODEL})...")
            # 設定 3 秒連線超時，120 秒讀取超時
            timeout_settings = httpx.Timeout(connect=3.0, read=120.0, write=120.0, pool=3.0)
            async with httpx.AsyncClient(timeout=timeout_settings) as client:
                resp = await client.post(
                    f"{OLLAMA_URL}/api/generate",
                    json={
                        "model":  OLLAMA_MODEL,
                        "prompt": prompt,
                        "stream": False,
                        "options": {"temperature": 0.3, "num_predict": 1500}
                    }
                )
                resp.raise_for_status()
                result = resp.json().get("response", "AI 回應解析失敗")
                print("[Mode: Local]")  # 本機成功時輸出 [Mode: Local]
                log.info("🟢 [Mode: Local] 本地 Ollama 推理成功")
                return result
        except (httpx.ConnectError, httpx.ConnectTimeout, httpx.TimeoutException, httpx.RequestError) as e:
            log.warning(f"⚠️ 本地 Ollama 連線失敗或超時 (3秒)：{e}。啟動雲端備援方案...")
            return await call_gemini_fallback(prompt)
        except Exception as e:
            log.warning(f"⚠️ 本地 Ollama 發生未知錯誤：{e}。啟動雲端備援方案...")
            return await call_gemini_fallback(prompt)


async def call_gemini_fallback(prompt: str) -> str:
    """
    雲端備援推論函數：使用 2026 最新官方 google-genai SDK 調用 gemini-2.5-flash
    """
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("環境變數 GEMINI_API_KEY 未設定，無法使用雲端備援模型")

        log.info("☁️ 呼叫雲端備援模型 gemini-2.5-flash...")
        # 初始化 2026 最新官方 google-genai SDK
        client = genai.Client(api_key=api_key)
        
        # 使用 aio 進行非同步呼叫，避免阻塞
        response = await client.aio.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
        )
        print("[Mode: Cloud]")  # 切換雲端成功時輸出 [Mode: Cloud]
        log.info("🟢 [Mode: Cloud] 雲端 Gemini 2.5 Flash 推理成功")
        return response.text
    except Exception as e:
        log.error(f"❌ 雲端備援推理失敗：{e}")
        return f"[AI 暫時離線] 錯誤訊息：地端與雲端模型皆無法使用。最後錯誤：{str(e)}"


# [DEMO FEATURE] 強化 Agent 行動模組以支援三大情境與個別行為區分
async def agent_action_notify(payload_log: str, risk_level: int) -> str:
    """
    Agent 主動通知與行動模組：根據不同異常情境執行對應的模擬防禦或通報行動
    """
    log_lower = payload_log.lower()
    
    # 情境二：資安威脅
    if any(k in log_lower for k in ["資安", "登入", "海外", "login", "ip", "匯出", "download", "leak"]):
        action_msg = "[Agent Action] 模擬發送資安告警信給 CISO，並已對來源 IP 啟動防禦封鎖阻斷"
        print(action_msg)  # 輸出至終端機
        return action_msg
        
    # 情境三：碳排超標
    elif any(k in log_lower for k in ["碳排", "14064", "電力", "度", "carbon"]):
        action_msg = "[Agent Action] 模擬計算碳排強度，生成改善報告並指派給永續發展委員會"
        print(action_msg)  # 輸出至終端機
        return action_msg
        
    # 情境一 / 預設：生管良率或高風險異常
    elif risk_level >= 2 or any(k in log_lower for k in ["良率", "smt", "生管", "sop"]):
        action_msg = "[Agent Action] 模擬發送緊急通知信給生管主管，並通知 SMT 區線長處置"
        print(action_msg)  # 輸出至終端機
        return action_msg
        
    return ""


# [FEATURE] HITL 審核機制
async def execute_agent_action(action_msg: str, approver: str):
    """
    執行核准後的實際 Agent 行動
    """
    log_msg = f"[HITL Action Execution] 審核人 {approver} 已核准，正在執行動作：{action_msg}"
    print(log_msg)  # 輸出至終端機
    log.info(log_msg)


# [FEATURE] 匯出正式報告功能
def build_docx_report(task_id: str, task_data: dict) -> io.BytesIO:
    doc = Document()
    
    # 設置標題 Style
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《企業數位韌性 - 國際標準合規與自主稽核改善報告》")
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 64, 175) # 深藍色 (Dark Blue)
    
    # 1. 基本資訊
    doc.add_heading("【一、基本資訊】", level=2)
    p_info = doc.add_paragraph()
    p_info.add_run("• 任務 ID: ").bold = True
    p_info.add_run(f"{task_id}\n")
    
    created_at = task_data.get("created_at", "")
    p_info.add_run("• 生成時間: ").bold = True
    p_info.add_run(f"{created_at}\n")
    
    risk_assessment = task_data.get("risk_assessment", {})
    risk_lbl = risk_assessment.get("risk_label", "Level-1 輕微")
    p_info.add_run("• 稽核項目與風險等級: ").bold = True
    p_info.add_run(f"{risk_lbl}\n")
    
    # 2. 異常數據摘要
    doc.add_heading("【二、異常數據摘要】", level=2)
    anon_data = task_data.get("anonymized_data", {})
    p_anon = doc.add_paragraph()
    p_anon.add_run("• 操作人員 ID (去識別化): ").bold = True
    p_anon.add_run(f"{anon_data.get('user_id', 'N/A')}\n")
    p_anon.add_run("• 部門 / 班別: ").bold = True
    p_anon.add_run(f"{anon_data.get('department', 'N/A')} / {anon_data.get('shift', 'N/A')}\n")
    p_anon.add_run("• 碳排放強度: ").bold = True
    p_anon.add_run(f"{anon_data.get('carbon_intensity', 0.0):.2f} kgCO₂e/unit\n")
    p_anon.add_run("• 當班良率: ").bold = True
    p_anon.add_run(f"{anon_data.get('yield_rate', 0.0):.1f}%\n" if anon_data.get('yield_rate') is not None else "N/A\n")
    p_anon.add_run("• 異常日誌: ").bold = True
    p_anon.add_run(f"{anon_data.get('event_log', 'N/A')}")
    
    # 3. AI 專家診斷與防呆建議
    doc.add_heading("【三、AI 專家診斷與防呆建議 (CAPA)】", level=2)
    capa_report = task_data.get("ai_capa_report", "")
    for line in capa_report.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    # 4. 審計與簽核軌跡
    doc.add_heading("【四、審計與簽核軌跡】", level=2)
    p_sign = doc.add_paragraph()
    p_sign.add_run("• 簽核狀態: ").bold = True
    p_sign.add_run("人類主管已核准執行 (Approved & Executed)\n")
    p_sign.add_run("• 核准人 (Approver): ").bold = True
    p_sign.add_run(f"{task_data.get('approved_by', 'N/A')}\n")
    p_sign.add_run("• 核准時間: ").bold = True
    p_sign.add_run(f"{task_data.get('approved_at', 'N/A')}\n")
    p_sign.add_run("• 執行的 Agent 動作: ").bold = True
    p_sign.add_run(f"{task_data.get('agent_notification', 'N/A')}\n")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


async def run_audit(task_id: str, payload: AuditPayload):
    """
    完整稽核流程（非同步背景執行）
    步驟：去識別化 → 風險評級 → RAG 檢索 → AI 推理 → 報告存儲
    """
    log.info(f"🔍 [{task_id}] 稽核任務啟動")
    report_store[task_id]["status"] = "processing"

    try:
        # Step 1: 去識別化
        safe_data = anonymize(payload)
        log.info(f"🔒 [{task_id}] 去識別化完成")

        # Step 2: 風險評級
        risk = risk_assessment(
            payload.carbon_intensity,
            payload.event_log,
            payload.production_yield
        )
        log.info(f"⚠️  [{task_id}] 風險評級：{risk['risk_label']}")

        # Step 3: RAG 向量檢索
        query = (
            f"碳排放 {payload.carbon_intensity} 超標 "
            f"資安事件 {payload.event_log[:50]} "
            f"CAPA 矯正預防措施"
        )
        context = retrieve_context(query)
        log.info(f"📚 [{task_id}] 向量檢索完成，上下文長度：{len(context)} 字元")

        # Step 4: 組裝 Prompt
        # [Step 5 ESG] 動態偵測 RAG 上下文是否包含 ESG/碳排相關內容，以決定是否啟用 ESG 輸出區塊
        _esg_keywords = ["碳排", "14064", "範疇", "溫室氣體", "能耗", "排放係數", "數據品質", "盤查", "carbon", "scope", "ghg", "emission"]
        _esg_context_detected = any(kw in context.lower() for kw in _esg_keywords)
        _carbon_is_high = payload.carbon_intensity > 1.0  # 碳排強度是否超標

        # [Step 5 ESG] 根據是否為 ESG 相關情境，動態組裝額外的 ESG 輸出指令
        _esg_prompt_addon = ""
        if _esg_context_detected or _carbon_is_high:
            _esg_prompt_addon = f"""
### 【ESG 盤查與減碳對策】（ISO 14064-1 主導稽核員專項輸出，必須填寫）
> 本事件涉及碳排放或能耗數據異常，依 ISO 14064-1 要求輸出以下盤查指引：

#### 排放邊界分析
（說明本事件影響的排放範疇：範疇一直接排放 / 範疇二電力間接排放，並說明受影響的廠區邊界）

#### 活動數據完整性查核
（確認能耗數據是否完整，缺漏件數、傳輸中斷是否已修復，是否符合 ISO 14064-1 §7.4 DQM 要求）

#### 排放係數合規性確認
（確認本期使用的電網排放係數是否為最新版本；燃料排放係數誤差是否在 ±10% 以內）

#### ESG 減碳行動建議（短／中／長期）
（依矯正 → 預防 → 系統性改善三層架構，提出具體可執行的減碳行動計畫）

#### 引用 ISO 14064-1 條文
（明確引用 §5.2/§5.3/§5.4/§7.4/§8/§9 等相關條文章節）
"""

        prompt = f"""
你是一位同時具備以下專業資格的企業合規顧問與 ESG 主導稽核員：
  - ISO 14064-1 溫室氣體盤查主導稽核員（GHG Lead Auditor）：專精組織碳盤查、範疇一/二排放管理、數據品質管理（DQM）與減碳 CAPA
  - ISO 27001 資訊安全管理主導稽核員：擅長資安事件響應、數位鑑識與稽核軌跡管理
  - 製造業生產管制工程師：熟悉 SMT 製程、良率管理、SOP 合規稽核

請根據以下資訊，產出一份**繁體中文**的標準 CAPA 矯正預防報告。
當事件涉及碳排放或能耗異常時，**必須**包含獨立的【ESG 盤查與減碳對策】區塊。

## 企業合規知識庫（ISO 條文與 SOP 摘要）
{context}

## 當前異常事件資料（已去識別化）
- 操作人員：{safe_data['user_id']}（{safe_data['department']}，{safe_data['shift']}班）
- 碳排放強度：{safe_data['carbon_intensity']:.2f} kgCO₂e/unit{"  ⚠️ 超標！合規上限 1.0" if _carbon_is_high else "  ✅ 合規"}
- 設備編號：{safe_data.get('equipment_id', 'N/A')}
- 良率：{safe_data.get('yield_rate', 'N/A')}%
- 資安日誌：{safe_data['event_log']}

## 風險評級結果
- 等級：{risk['risk_label']}
- 判定原因：{'; '.join(risk['reasons'])}
- 要求回應時限：{risk['response_time']}

## 請輸出以下結構的 CAPA 報告（每項須具體、可執行）

### 1. 事件摘要
（簡述事件發生背景，使用 5W1H 格式）

### 2. 直接原因分析
（導致本次異常的直接觸發因素）

### 3. 根本原因分析（5-Why）
（至少進行 3 層 Why 追問）

### 4. 矯正措施（Corrective Action）
| 項次 | 措施內容 | 負責人層級 | 完成時限 |
|------|----------|------------|----------|
（列出 3-5 項具體可執行的矯正行動）

### 5. 預防措施（Preventive Action）
（列出 2-3 項防止復發的系統性改善方案）

### 6. 效果驗證指標
（說明如何量化驗證改善成效）

### 7. 引用法規條文
（明確引用相關 ISO 條文與 SOP 章節）
{_esg_prompt_addon}"""

        # Step 5: AI 推理
        log.info(f"🤖 [{task_id}] 啟動 AI 雙引擎推理...")
        ai_response = await call_ai_inference(prompt)

        # [FEATURE] 呼叫 Agent 主動通知機制
        notify_action = await agent_action_notify(payload.event_log, risk["risk_level"])

        # [FEATURE] HITL 審核機制：暫存擬定的 Agent 行動，原本狀態為 completed，現在改為 pending_approval
        # Step 6: 儲存報告（記憶體 + SQLite 持久化）
        report_store[task_id].update({
            "status":             "pending_approval",
            "anonymized_data":    safe_data,
            "risk_assessment":    risk,
            "ai_capa_report":     ai_response,
            "proposed_action":    notify_action,  # 暫存擬定行動
            "context_used":       context[:500] + "..."
        })
        # [Step 4] 持久化至 SQLite
        save_report(task_id, report_store[task_id])
        log.info(f"✅ [{task_id}] 稽核完成，報告已生成並持久化，等待審核")

    except Exception as e:
        log.error(f"❌ [{task_id}] 稽核任務失敗：{e}")
        report_store[task_id].update({
            "status": "failed",
            "error":  str(e)
        })
        # [Step 4] 失敗狀態也持久化
        save_report(task_id, report_store[task_id])


# ╔══════════════════════════════════════════════════════════╗
# ║  API 路由                                                ║
# ╚══════════════════════════════════════════════════════════╝

# [FEATURE] RBAC 權限控制：核發 Token 端點
@app.post("/token", tags=["身份驗證"])
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    """核發 JWT Token"""
    user = USERS.get(form_data.username)
    if not user or user["password"] != form_data.password:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="帳號或密碼錯誤",
            headers={"WWW-Authenticate": "Bearer"},
        )
    token_data = {"sub": form_data.username, "role": user["role"]}
    access_token = jwt.encode(token_data, SECRET_KEY, algorithm=ALGORITHM)
    return {"access_token": access_token, "token_type": "bearer", "role": user["role"]}


@app.post("/api/login", tags=["身份驗證"])
async def api_login(payload: LoginRequest):
    """自訂登入 API，回傳假 JWT Token 與 Role"""
    user = USERS.get(payload.username)
    if not user or user["password"] != payload.password:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="帳號或密碼錯誤"
        )
    token_data = {"sub": payload.username, "role": user["role"]}
    token = jwt.encode(token_data, SECRET_KEY, algorithm=ALGORITHM)
    return {"token": token, "role": user["role"]}


@app.get("/api/health/ollama", tags=["健康檢查"])
async def health_ollama():
    """公開 (Public) Ollama 健康檢查端點，使用 requests 測試 2 秒 timeout"""
    try:
        # 使用 requests 設定 2 秒 timeout 去 ping http://localhost:11434/
        r = requests.get("http://localhost:11434/", timeout=2.0)
        if r.status_code == 200:
            return {"status": "online"}
        return {"status": "offline"}
    except Exception:
        return {"status": "offline"}


@app.post("/api/audit/approve_sec", tags=["報告審核"])
async def approve_security_alert(current_user: User = Depends(get_current_user)):
    """核准資安警報 API，限 admin 或 ciso 權限"""
    if current_user.role not in ["admin", "ciso"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="權限不足，只有 admin 或 ciso 可以核准資安警報"
        )
    
    # 寫入 audit.db 稽核軌跡
    from database import log_audit_event
    timestamp_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_audit_event(
        username=current_user.username,
        role=current_user.role,
        action_type="Approval",
        action_details=f"[{timestamp_str}] {current_user.username.upper()} 已核准 SEC-889 任務"
    )
    
    return {"status": "success", "message": "資安警報已成功核准"}


class AdviceRequest(BaseModel):
    prompt: str


# ── 異步任務佇列管理 (Task Queue) ──────────────────────────────
ai_task_queue = asyncio.Queue()
ai_tasks_registry = {}

async def ai_inference_worker():
    """背景 AI 佇列工作者，Concurrency = 1，確保地端 Ollama 同時間只有一個任務在執行。"""
    import time
    import random
    from database import log_audit_event
    from guardrails import LLMGuardrail
    from rag_service import query_knowledge

    log.info("[Queue Worker] 背景任務排隊 Worker 已啟動，等待任務中...")
    while True:
        try:
            task_data = await ai_task_queue.get()
            task_id = task_data["task_id"]
            prompt = task_data["prompt"]
            masked_prompt = task_data["masked_prompt"]
            mask_vault = task_data["mask_vault"]
            username = task_data["username"]
            role = task_data["role"]
            start_time = task_data["start_time"]

            ai_tasks_registry[task_id]["status"] = "RUNNING"
            log.info(f"⚙️ [Queue Worker] 任務 {task_id} 開始推理...")

            # 1. 呼叫 RAG 知識庫檢索
            context = query_knowledge(masked_prompt, n_results=3)

            # 2. 模擬地端 AI 推理過程 (改用 asyncio.sleep 避免阻塞 event loop)
            await asyncio.sleep(random.uniform(1.5, 3.0))

            prompt_lower = masked_prompt.lower()
            if "資安" in prompt_lower or "事件" in prompt_lower or "下載" in prompt_lower or "sec" in prompt_lower:
                ai_res_raw = (
                    f"【地端 AI (RAG 增強建議)】\n"
                    f"根據檢索到的 ISO 27001 A.8.16 參考條文，針對大量下載行為，建議採取以下對策：\n"
                    f"1. 對涉事 IP 進行網絡端口阻斷；\n"
                    f"2. 隔離涉事主體 (已遮蔽代碼: {masked_prompt})；\n"
                    f"3. 立即重置該帳號的 SSO 登入 Token，並在 2 小時內向 CISO 進行事故通報。"
                )
            elif "碳" in prompt_lower or "carbon" in prompt_lower or "14064" in prompt_lower:
                ai_res_raw = (
                    f"【地端 AI (RAG 增強建議)】\n"
                    f"根據檢索到的 ISO 14064-1 參考條文，碳強度指標超標矯正措施：\n"
                    f"1. 立即啟動 Level-1 / Level-2 矯正改善程序；\n"
                    f"2. 自動指派碳強度異常檢討單至永續發展委員會，通報廠長召開緊急減碳會議 (已遮蔽代碼: {masked_prompt})。"
                )
            else:
                ai_res_raw = (
                    f"【地端 AI (RAG 增強建議)】\n"
                    f"根據檢索到的 SOP-PROD-001 製造規範，良率低於目標 85.0% 矯正對策：\n"
                    f"1. 製程組長與品保工程師應於 4 小時內草擬 CAPA 報告；\n"
                    f"2. 立即排查貼片機吸嘴磨損情況，並校正回焊爐加熱區溫度設定。"
                )

            latency = int((time.time() - start_time) * 1000)

            # 3. 呼叫輸出護欄還原個資 (Unmask)
            ai_res_unmasked = LLMGuardrail.unmask_response(ai_res_raw, mask_vault)

            # 4. 記錄到資料庫日誌 (儲存去識別化後的 masked 資訊，符合安全原則)
            log_audit_event(
                username=username,
                role=role,
                action_type="LLM_Generation",
                action_details="向地端大腦 (Ollama) 請求 AI 改善對策建議 (已通過 RAG 知識檢索與 I/O 護欄)",
                prompt=masked_prompt,
                ai_response=ai_res_raw,
                latency_ms=latency
            )

            # 5. 更新狀態表為 SUCCESS
            ai_tasks_registry[task_id].update({
                "status": "SUCCESS",
                "advice": ai_res_unmasked,
                "latency_ms": latency
            })
            log.info(f"✅ [Queue Worker] 任務 {task_id} 處理完成")

        except Exception as e:
            log.error(f"❌ [Queue Worker] 任務處理失敗: {e}")
            if "task_id" in locals():
                ai_tasks_registry[task_id].update({
                    "status": "FAILED",
                    "error": str(e)
                })
        finally:
            ai_task_queue.task_done()


@app.on_event("startup")
async def startup_event():
    # 啟動背景 Worker，執行 Concurrency = 1 的 AI 推理排隊
    asyncio.create_task(ai_inference_worker())

    # [Step 4] 重啟恢復機制：從 SQLite 還原 pending_approval / processing 任務至記憶體
    try:
        recovered = get_reports_for_recovery(limit=50)
        for report_data in recovered:
            tid = report_data.get("task_id")
            if tid and tid not in report_store:
                report_store[tid] = report_data
                log.info(f"♻️  [Recovery] 已從 DB 還原任務：{tid} ({report_data.get('status')})")
        if recovered:
            log.info(f"♻️  [Recovery] 共還原 {len(recovered)} 筆待處理報告至記憶體快取")
    except Exception as e:
        log.warning(f"⚠️  [Recovery] 啟動恢復失敗（非致命）：{e}")


@app.post("/api/ai/generate_advice", tags=["AI 諮詢"])
async def generate_advice(payload: AdviceRequest, current_user: User = Depends(get_current_user)):
    """向 AI 請求分析建議，需驗證 Token。包含輸入護欄與 Prompt Injection 阻攔，並排入背景任務佇列"""
    import time
    from database import log_audit_event
    from guardrails import LLMGuardrail

    start_time = time.time()

    # 1. 呼叫輸入護欄 (PII 遮蔽與 Prompt Injection 偵測)
    masked_prompt, mask_vault, is_blocked, block_reason = LLMGuardrail.mask_prompt(payload.prompt)

    # 2. 判斷是否攔截 Prompt Injection
    if is_blocked:
        # 寫入審計日誌 (即使被阻擋，這也是一項關鍵安全事件)
        log_audit_event(
            username=current_user.username,
            role=current_user.role,
            action_type="Security_Block",
            action_details=f"安全中樞成功攔截 Prompt Injection 攻擊事件。原因：{block_reason}",
            prompt=payload.prompt,
            ai_response=block_reason,
            latency_ms=0
        )
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=block_reason
        )

    # 3. 建立並加入佇列
    task_id = str(uuid.uuid4())
    ai_tasks_registry[task_id] = {
        "status": "PENDING",
        "prompt": payload.prompt,
        "masked_prompt": masked_prompt,
        "mask_vault": mask_vault,
        "username": current_user.username,
        "role": current_user.role,
        "start_time": start_time,
        "advice": None,
        "latency_ms": None
    }

    await ai_task_queue.put({
        "task_id": task_id,
        "prompt": payload.prompt,
        "masked_prompt": masked_prompt,
        "mask_vault": mask_vault,
        "username": current_user.username,
        "role": current_user.role,
        "start_time": start_time
    })

    log.info(f"📥 任務 {task_id} 已成功加入佇列，狀態為 PENDING")
    return {"task_id": task_id, "status": "PENDING"}


@app.get("/api/ai/task_status/{task_id}", tags=["AI 諮詢"])
async def get_task_status(task_id: str, current_user: User = Depends(get_current_user)):
    """查詢任務處理狀態與結果，需驗證 Token"""
    if task_id not in ai_tasks_registry:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="找不到指定的任務 ID"
        )

    task = ai_tasks_registry[task_id]
    result = {
        "task_id": task_id,
        "status": task["status"]
    }
    if task["status"] == "SUCCESS":
        result["advice"] = task["advice"]
        result["latency_ms"] = task["latency_ms"]
    elif task["status"] == "FAILED":
        result["error"] = task.get("error", "任務執行失敗")

    return result


@app.get("/api/audit/logs", tags=["審計日誌"])
async def get_logs(current_user: User = Depends(get_current_user)):
    """查詢最新 50 筆稽核軌跡日誌，僅限 admin 或 ciso 權限"""
    if current_user.role not in ["admin", "ciso"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="權限不足，只有 admin 或 ciso 可以查看稽核軌跡日誌"
        )
    from database import get_audit_logs
    return get_audit_logs(limit=50)


@app.get("/api/reports/generate", tags=["報表匯出"])
async def generate_weekly_report(current_user: User = Depends(get_current_user)):
    """
    撈取近 7 天的稽核軌跡日誌，呼叫 AI 進行合規摘要後，生成專業 Word 稽核報告並回傳下載 (Step 5)。
    """
    if current_user.role not in ["admin", "ciso"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="權限不足，只有 admin 或 ciso 可以生成並下載稽核報告"
        )
        
    from database import get_audit_logs_since
    from report_service import generate_audit_word_report
    import io
    
    # 1. 撈取近 7 天日誌
    logs = get_audit_logs_since(days=7)
    
    # 2. AI 摘要分析 (防呆：無資料時免除 AI 呼叫)
    if not logs:
        ai_summary = "本週系統運作良好，無任何安全或環境合規異常事件記錄。"
    else:
        # 將日誌轉換為精簡字串以組裝 Prompt
        log_lines = []
        for r in logs:
            ts = r.get("timestamp", "")
            user = r.get("username", "System")
            role = r.get("role", "SYSTEM")
            action = r.get("action_type", "")
            details = r.get("action_details", "")
            p = r.get("prompt", "")
            res = r.get("ai_response", "")
            
            line = f"時間: {ts} | 使用者: {user}({role}) | 類型: {action} | 詳情: {details}"
            if p:
                line += f" | 輸入: {p}"
            if res:
                # 截短建議，防止超過模型 context 長度
                line += f" | AI建議: {res[:80]}..."
            log_lines.append(line)
            
        logs_summary = "\n".join(log_lines)
        
        prompt = f"""
你是一位資深的主導稽核員（Lead Auditor），專長為 ISO 27001 資訊安全管理系統與 ISO 14064-1 溫室氣體盤查規範。
請根據以下過去 7 天內系統的稽核軌跡日誌（包含使用者操作、AI 生成對策、以及安全防禦攔截），撰寫一份約 300 字的「本週整體系統合規與風險狀態總結（Executive Summary）」。

要求：
1. 口吻與文字需非常專業、嚴謹，符合正式 ISO 查核報告的正式格式與語彙（如「建議採取矯正預防措施 (CAPA)」、「未發現顯著不符合事項 (Non-conformity)」、「安全控制措施運作有效」等）。
2. 扼要摘要本週的重要事件（如：有哪些安全警報、審核批准活動、或是 ESG 碳排放/良率等指標異常）。
3. 給出客觀的合規性風險評估，並以繁體中文撰寫。

【稽核軌跡日誌如下】：
{logs_summary}
"""
        try:
            # 呼叫雙引擎 AI 進行摘要 (內含 Ollama / Gemini Fallback 機制)
            ai_summary = await call_ai_inference(prompt)
        except Exception as e:
            ai_summary = f"（AI 摘要生成出錯：{str(e)}。已直接載入預設報告。）\n本週系統有相關稽核日誌運作，安全與環境合規控制措施運作中。"
            
    # 3. 呼叫 Word 生成模組
    try:
        report_bytes = generate_audit_word_report(logs, ai_summary)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Word 報告檔案生成失敗：{str(e)}"
        )
        
    # 4. 回傳 StreamingResponse
    today_filename = datetime.now().strftime("%Y%m%d")
    headers = {
        "Content-Disposition": f"attachment; filename=Audit_Report_{today_filename}.docx",
        "Access-Control-Expose-Headers": "Content-Disposition"
    }
    return StreamingResponse(
        io.BytesIO(report_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@app.get("/", tags=["健康檢查"])
async def root():
    return {
        "system":  "企業數位韌性 AI 稽核系統",
        "version": "2.0.0",
        "status":  "online",
        "model":   OLLAMA_MODEL,
        "time":    datetime.now().isoformat()
    }


@app.get("/health", tags=["健康檢查"])
async def health():
    """檢查各子系統狀態"""
    checks = {}

    # 檢查 Ollama
    try:
        async with httpx.AsyncClient(timeout=0.5) as client:
            r = await client.get(f"{OLLAMA_URL}/api/tags")
            checks["ollama"] = "online" if r.status_code == 200 else "degraded"
    except Exception:
        checks["ollama"] = "offline"

    # 檢查 ChromaDB
    try:
        client = chromadb.PersistentClient(path=DB_PATH)
        cols = client.list_collections()
        checks["chromadb"] = f"online ({len(cols)} collections)"
    except Exception:
        checks["chromadb"] = "offline"

    checks["api"]         = "online"
    checks["active_tasks"] = len([v for v in report_store.values() if v.get("status") == "processing"])

    return checks


@app.get("/api/health/detailed", tags=["健康檢查"])
async def health_detailed():
    """[Step 4] 詳細健康狀態端點，含 DB 統計、任務佇列深度與 AI Worker 狀態"""
    result = {}

    # Ollama 狀態
    try:
        async with httpx.AsyncClient(timeout=1.0) as client:
            r = await client.get(f"{OLLAMA_URL}/api/tags")
            result["ollama"] = "online" if r.status_code == 200 else "degraded"
    except Exception:
        result["ollama"] = "offline"

    # ChromaDB 狀態
    try:
        chroma_client = chromadb.PersistentClient(path=DB_PATH)
        cols = chroma_client.list_collections()
        result["chromadb"] = f"online ({len(cols)} collections)"
    except Exception:
        result["chromadb"] = "offline"

    # SQLite 資料庫統計
    try:
        db_stats = get_db_stats()
        result["sqlite"] = "online"
        result["db_stats"] = db_stats
    except Exception as e:
        result["sqlite"] = f"error: {e}"
        result["db_stats"] = {}

    # 任務佇列狀態
    result["task_queue_depth"]    = ai_task_queue.qsize()
    result["memory_report_count"] = len(report_store)
    result["active_tasks"]        = len([v for v in report_store.values() if v.get("status") == "processing"])
    result["pending_approval"]    = len([v for v in report_store.values() if v.get("status") == "pending_approval"])
    result["api"]                 = "online"
    result["timestamp"]           = datetime.now().isoformat()

    return result


# [FEATURE] 導入非同步狀態追蹤，設定初始狀態為 processing
@app.post("/api/v1/trigger_audit", tags=["稽核觸發"])
async def trigger_audit(payload: AuditPayload, background_tasks: BackgroundTasks):
    """
    接收異常數據，非同步觸發 AI 稽核流程
    （適用於 Power BI Webhook 或自動化排程呼叫）
    """
    task_id = str(uuid.uuid4())[:8]
    report_store[task_id] = {
        "task_id":    task_id,
        "status":     "processing",  # 狀態改為 processing
        "created_at": datetime.now().isoformat(),
        "input": {
            "carbon_intensity": payload.carbon_intensity,
            "department":       payload.department,
            "shift":            payload.shift,
        }
    }

    background_tasks.add_task(run_audit, task_id, payload)

    log.info(f"📥 新稽核任務已接收：{task_id}")
    return {
        "status":    "processing",  # 返回狀態設為 processing
        "task_id":   task_id,
        "message":   "稽核任務已非同步觸發，請使用 task_id 查詢結果",
        "query_url": f"/api/task/{task_id}"  # 指向新的狀態查詢 API
    }


# [FEATURE] 新增狀態查詢 API
@app.get("/api/task/{task_id}", tags=["報告查詢"])
async def get_task_report(task_id: str):
    """狀態查詢 API：回傳任務當前的處理狀態或最終結果（Memory-First，次取 SQLite）"""
    # 優先從記憶體快取讀取（速度快）
    if task_id in report_store:
        return report_store[task_id]
    # Fallback：從 SQLite 讀取（跨重啟資料恢復）
    db_record = get_report(task_id)
    if db_record:
        report_store[task_id] = db_record  # 回填快取
        return db_record
    raise HTTPException(status_code=404, detail=f"任務 {task_id} 不存在")


# [FEATURE] HITL 審核機制：核准並執行端點
# [FEATURE] RBAC 權限控制：引入 JWT 驗證與權限控制
@app.post("/api/task/{task_id}/approve", tags=["報告審核"])
async def approve_task(task_id: str, payload: ApprovePayload, current_user: User = Depends(get_current_user)):
    """核准任務並執行實際的 Agent 動作"""
    if task_id not in report_store:
        raise HTTPException(status_code=404, detail=f"任務 {task_id} 不存在")
    
    task_data = report_store[task_id]
    if task_data.get("status") != "pending_approval":
        raise HTTPException(status_code=400, detail=f"任務狀態為 {task_data.get('status')}，非待審核狀態")
    
    # [FEATURE] RBAC 權限控制：根據任務類型與登入者角色進行權限校驗
    log_text = task_data.get("anonymized_data", {}).get("event_log", "").lower()
    is_security = any(k in log_text for k in ["資安", "登入", "海外", "login", "ip", "匯出", "download", "leak", "siem"])
    is_yield = any(k in log_text for k in ["良率", "smt", "生管", "sop", "生產"])
    
    if current_user.role == "admin":
        pass
    elif current_user.role == "ciso":
        if is_yield:
            raise HTTPException(status_code=403, detail="您沒有權限核准生產良率稽核任務")
    elif current_user.role == "qms_manager":
        if is_security:
            raise HTTPException(status_code=403, detail="您沒有權限核准資訊安全稽核任務")
    else:
        raise HTTPException(status_code=403, detail="未授權的角色")
        
    proposed_action = task_data.get("proposed_action", "")
    
    # 執行實際的 Agent 動作
    await execute_agent_action(proposed_action, payload.approver)
    
    # 更新狀態為 completed，並記錄審核人與審核時間
    task_data.update({
        "status": "completed",
        "completed_at": datetime.now().isoformat(),
        "approved_by": payload.approver,
        "approved_at": datetime.now().isoformat(),
        "agent_notification": f"[已執行] 由 {payload.approver} ({current_user.role}) 於 {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 核准並執行：{proposed_action}"
    })
    # [Step 4] 持久化核准結果至 SQLite
    save_report(task_id, task_data)
    return task_data


# [FEATURE] HITL 審核機制：退回重擬端點
# [FEATURE] RBAC 權限控制：引入 JWT 驗證與權限控制
@app.post("/api/task/{task_id}/reject", tags=["報告審核"])
async def reject_task(task_id: str, payload: RejectPayload, current_user: User = Depends(get_current_user)):
    """退回任務並填寫退回原因"""
    if task_id not in report_store:
        raise HTTPException(status_code=404, detail=f"任務 {task_id} 不存在")
    
    task_data = report_store[task_id]
    if task_data.get("status") != "pending_approval":
        raise HTTPException(status_code=400, detail=f"任務狀態為 {task_data.get('status')}，非待審核狀態")
        
    # [FEATURE] RBAC 權限控制：根據任務類型與登入者角色進行權限校驗
    log_text = task_data.get("anonymized_data", {}).get("event_log", "").lower()
    is_security = any(k in log_text for k in ["資安", "登入", "海外", "login", "ip", "匯出", "download", "leak", "siem"])
    is_yield = any(k in log_text for k in ["良率", "smt", "生管", "sop", "生產"])
    
    if current_user.role == "admin":
        pass
    elif current_user.role == "ciso":
        if is_yield:
            raise HTTPException(status_code=403, detail="您沒有權限退回生產良率稽核任務")
    elif current_user.role == "qms_manager":
        if is_security:
            raise HTTPException(status_code=403, detail="您沒有權限退回資訊安全稽核任務")
    else:
        raise HTTPException(status_code=403, detail="未授權的角色")
    
    # 更新狀態為 rejected，並記錄退回人與退回原因
    task_data.update({
        "status": "rejected",
        "rejected_at": datetime.now().isoformat(),
        "rejected_by": payload.rejected_by,
        "reject_reason": payload.reason,
        "agent_notification": f"[已退回] 由 {payload.rejected_by} ({current_user.role}) 於 {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} 退回，原因：{payload.reason}"
    })
    # [Step 4] 持久化退回結果至 SQLite
    save_report(task_id, task_data)
    return task_data


# [DEMO FEATURE] 新增情境腳本展示數據查詢 API
@app.get("/api/v1/scenarios", tags=["情境模擬"])
async def get_scenarios():
    """情境腳本展示模組：回傳定義好的三個 Mock 情境數據"""
    return DEMO_SCENARIOS


@app.get("/api/v1/report/{task_id}", tags=["報告查詢"])
async def get_single_report(task_id: str):
    """[Step 4] 查詢指定任務的完整稽核報告（Memory-First，次取 SQLite）"""
    if task_id in report_store:
        return report_store[task_id]
    db_record = get_report(task_id)
    if db_record:
        report_store[task_id] = db_record
        return db_record
    raise HTTPException(status_code=404, detail=f"任務 {task_id} 不存在")


# [FEATURE] 匯出正式報告功能
# [FEATURE] RBAC 權限控制：引入 JWT 驗證與權限控制
@app.get("/api/task/{task_id}/export", tags=["報告匯出"])
async def export_task_report(task_id: str, current_user: User = Depends(get_current_user)):
    """將已核准完成的報告匯出為正式 Word (DOCX) 文件"""
    if task_id not in report_store:
        raise HTTPException(status_code=404, detail=f"任務 {task_id} 不存在")
    
    task_data = report_store[task_id]
    if task_data.get("status") != "completed":
        raise HTTPException(status_code=400, detail="任務尚未核准完成，無法匯出正式報告")
        
    # [FEATURE] RBAC 權限控制：根據任務類型與登入者角色進行下載權限校驗
    log_text = task_data.get("anonymized_data", {}).get("event_log", "").lower()
    is_security = any(k in log_text for k in ["資安", "登入", "海外", "login", "ip", "匯出", "download", "leak", "siem"])
    
    if current_user.role == "qms_manager" and is_security:
        raise HTTPException(status_code=403, detail="您沒有權限下載資訊安全稽核報告")
    
    # 產生 DOCX 檔案流
    file_stream = build_docx_report(task_id, task_data)
    
    headers = {
        "Content-Disposition": f"attachment; filename=ISO_Audit_Report_{task_id}.docx"
    }
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        headers=headers
    )


@app.get("/api/v1/reports", tags=["報告查詢"])
async def list_all_reports():
    """[Step 4] 列出所有稽核任務摘要（從 SQLite 讀取，保證跨重啟持久）"""
    db_reports = list_reports(limit=200)
    # 以記憶體中最新狀態覆蓋 DB 快照（確保 in-progress 任務狀態最新）
    mem_lookup = {
        tid: {
            "task_id":    tid,
            "status":     data.get("status"),
            "created_at": data.get("created_at"),
            "risk_level": data.get("risk_assessment", {}).get("risk_level", 1),
            "risk_label": data.get("risk_assessment", {}).get("risk_label", "待評估"),
        }
        for tid, data in report_store.items()
    }
    result = []
    seen = set()
    for r in db_reports:
        tid = r["task_id"]
        seen.add(tid)
        if tid in mem_lookup:
            result.append(mem_lookup[tid])
        else:
            result.append(r)
    # 加入 DB 中尚未出現的記憶體任務（剛觸發還未寫 DB 的極短暫狀態）
    for tid, data in mem_lookup.items():
        if tid not in seen:
            result.append(data)
    # 依 created_at 降序排列
    result.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return result


@app.post("/api/v1/quick_risk", tags=["快速評估"])
async def quick_risk(payload: AuditPayload):
    """
    快速風險評估（同步回應，不啟動 AI 推理）
    適用於即時告警場景
    """
    risk = risk_assessment(
        payload.carbon_intensity,
        payload.event_log,
        payload.production_yield
    )
    return {
        "risk":    risk,
        "suggest": "建議立即觸發 /api/v1/trigger_audit 進行完整 AI 稽核" if risk["risk_level"] >= 2 else "持續監控中"
    }
